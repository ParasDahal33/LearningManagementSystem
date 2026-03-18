"""
parsers/matching_parser.py
Matching / drag-and-drop question parsing for all parser versions.
Handles table grid extraction, column scoring, pair extraction,
and instruction-table filtering.
"""

from __future__ import annotations

import hashlib
import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .util import (
    clean_text,
    is_red_hex,
    normalize_key,
    strip_q_prefix,
    v2_clean_text,
    v2_normalize_key,
    v3_clean_text,
    v3_normalize_key,
    v3_strip_q_prefix,
)
from .docx_extractor import (
    iter_block_items,
    paragraph_text_and_is_red,
    v3_iter_all_tables,
    v3_iter_tables_recursive,
)

# ---------------------------------------------------------------------------
# Shared matching stem regexes
# ---------------------------------------------------------------------------
MATCHING_STEM_RE = re.compile(
    r"\b(complete\s+the\s+table|drag(?:ging)?\s+and\s+drop(?:ping)?|drag\s+and\s+drop|"
    r"match\s+each|match\s+the\s+following|match\s+.*\s+to\s+the\s+correct|select\s+one.*for\s+each)\b",
    re.IGNORECASE,
)
COMMAND_QUESTION_RE_V1 = re.compile(
    r"^(Illustrate|Critically\s+(?:assess|analyse|analyze|evaluate)|"
    r"Evaluate|Determine|Articulate|Prescribe|Analyse|Analyze|Review|Recommend)\b.+",
    re.IGNORECASE,
)

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def join_lines(lines: list[str]) -> str:
    parts = [clean_text(x) for x in (lines or []) if clean_text(x)]
    return "; ".join(parts).strip()


def table_to_grid(tbl: Table) -> list[list[list[str]]]:
    grid: list[list[list[str]]] = []
    for row in tbl.rows:
        r: list[list[str]] = []
        for cell in row.cells:
            lines = [clean_text(p.text) for p in cell.paragraphs if clean_text(p.text)]
            r.append(lines)
        grid.append(r)
    return grid


def table_fingerprint(grid: list[list[list[str]]]) -> str:
    rows = ["|".join(join_lines(c) for c in r) for r in grid]
    return normalize_key("||".join(rows))


def score_columns(grid, a: int, b: int) -> int:
    sc = 0
    for r in grid[1:]:
        if a >= len(r) or b >= len(r):
            continue
        if join_lines(r[a]) and join_lines(r[b]):
            sc += 1
    return sc


def pick_best_columns(grid):
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    best = None
    best_sc = 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = score_columns(grid, a, b)
            if sc > best_sc:
                best_sc = sc
                best = (a, b)
    return best if best_sc >= 2 else None


def extract_pairs(grid, left_col: int, right_col: int, start_row: int = 1) -> list[dict]:
    pairs = []
    for r in grid[start_row:]:
        if left_col >= len(r) or right_col >= len(r):
            continue
        left = join_lines(r[left_col])
        right = join_lines(r[right_col])
        if left and right:
            pairs.append({"left": left, "right": right})
    return pairs


def looks_like_matching_stem(t: str) -> bool:
    t2 = strip_q_prefix(clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors", "for students")):
        return False
    return bool(MATCHING_STEM_RE.search(t2))


# ===========================================================================
# V1 matching parser
# ===========================================================================
def parse_matching_questions_doc_order(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    recent: list[str] = []
    seen: set[str] = set()
    seq = 0

    def choose_stem() -> str | None:
        for t in reversed(recent[-50:]):
            if looks_like_matching_stem(t):
                return strip_q_prefix(t)
        return None

    for el in iter_block_items(doc):
        seq += 1
        if isinstance(el, Paragraph):
            t, _ = paragraph_text_and_is_red(el)
            t = clean_text(t)
            if t:
                recent.append(t)
            continue
        grid = table_to_grid(el)
        fp = table_fingerprint(grid)
        if fp in seen:
            continue
        seen.add(fp)
        cols = pick_best_columns(grid)
        if not cols:
            continue
        left_col, right_col = cols
        pairs = extract_pairs(grid, left_col, right_col, start_row=1)
        if len(pairs) < 2:
            continue
        stem = choose_stem()
        if not stem:
            header = grid[0] if grid else []
            hL = join_lines(header[left_col]) if header and left_col < len(header) else "Left"
            hR = join_lines(header[right_col]) if header and right_col < len(header) else "Right"
            stem = f"Match each '{hL}' to the correct '{hR}'."
        out.append({
            "question": stem,
            "pairs": pairs,
            "kind": "matching",
            "options": [],
            "correct": [],
            "multi": False,
            "_order": seq,
        })
        recent = []
    return out


# ===========================================================================
# V1-exact matching parser (richer heuristics)
# ===========================================================================
MATCHING_STEM_RE_V1_EXACT = re.compile(
    r"\b(complete\s+the\s+table|drag(?:ging)?\s+and\s+drop(?:ping)?|drag\s+and\s+drop|"
    r"match\s+each|match\s+the|match\s+.*\s+to\s+the\s+correct|select\s+one.*for\s+each)\b",
    re.IGNORECASE,
)
INSTRUCTION_TABLE_NOISE_RE = re.compile(
    r"\b(range\s+and\s+conditions?|decision-?making\s+rules?|rubric|pre-?approved\s+reasonable\s+adjustments?|"
    r"for\s+learners?|for\s+assessors?|instructions?|evidence|required|criteria|competent|nyc|submission|marking)\b",
    re.IGNORECASE,
)


def _looks_like_matching_stem_v1_exact(t: str) -> bool:
    t2 = strip_q_prefix(clean_text(t))
    if not t2 or t2.lower().startswith(("for learners", "for assessors")):
        return False
    if COMMAND_QUESTION_RE_V1.match(t2) or "which of the following" in t2.lower():
        return False
    return bool(MATCHING_STEM_RE_V1_EXACT.search(t2))


def _cell_lines_v1_exact(cell) -> list[str]:
    lines: list[str] = []
    for p in cell.paragraphs:
        t, _ = paragraph_text_and_is_red(p)
        t = clean_text(t)
        if not t:
            continue
        if "•" in t:
            lines.extend(x.strip() for x in t.split("•") if x.strip())
        else:
            lines.append(t)
    seen: set[str] = set()
    out: list[str] = []
    for x in lines:
        x = clean_text(x)
        k = normalize_key(x)
        if x and k not in seen:
            seen.add(k)
            out.append(x)
    return out


def _table_to_grid_v1_exact(table: Table) -> list[list[list[str]]]:
    return [[_cell_lines_v1_exact(c) for c in row.cells] for row in table.rows]


def _table_fingerprint_v1_exact(grid) -> str:
    flat = ["|".join(cell) for row in grid for cell in row]
    blob = "||".join([normalize_key(x) for x in flat if x])
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()


def _is_instruction_table_v1_exact(grid) -> bool:
    texts = [" ".join(cell) for row in grid for cell in row if cell]
    if not texts:
        return True
    first_row = " ".join(join_lines(c) for c in (grid[0] if grid else [])).lower()
    if "range and conditions" in first_row or "decision-making rules" in first_row or "pre-approved" in first_row:
        return True
    hits = sum(1 for t in texts if INSTRUCTION_TABLE_NOISE_RE.search(t))
    return (hits / max(1, len(texts))) >= 0.40


def _pair_is_valid_v1_exact(left: str, right: str) -> bool:
    if not left or not right:
        return False
    if normalize_key(left) == normalize_key(right):
        return False
    return len(left) <= 180 and len(right) <= 350


def _guess_header_skip_v1_exact(grid) -> int:
    if not grid or not grid[0]:
        return 0
    row0 = " ".join(join_lines(c) for c in grid[0] if c).strip().lower()
    header_words = ["definition", "term", "meaning", "word", "concept", "numbers", "number", "example", "type", "classification", "left", "right"]
    if any(w in row0 for w in header_words):
        return 1
    nonempty = [join_lines(c) for c in grid[0] if join_lines(c)]
    if nonempty and sum(len(x) <= 20 for x in nonempty) / len(nonempty) >= 0.8:
        return 1
    return 0


def _score_columns_v1_exact(grid, a: int, b: int) -> int:
    start = _guess_header_skip_v1_exact(grid)
    return sum(
        1
        for r in range(start, len(grid))
        if a < len(grid[r]) and b < len(grid[r]) and _pair_is_valid_v1_exact(join_lines(grid[r][a]), join_lines(grid[r][b]))
    )


def _pick_best_columns_v1_exact(grid) -> tuple[int, int] | None:
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    if max_cols < 2:
        return None
    best, best_score = None, 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = _score_columns_v1_exact(grid, a, b)
            if sc > best_score:
                best_score = sc
                best = (a, b)
    return best if best and best_score >= 2 else None


def _extract_pairs_v1_exact(grid, left_col: int, right_col: int, start_row: int = 0) -> list[dict]:
    pairs = []
    for r in range(start_row, len(grid)):
        if left_col >= len(grid[r]) or right_col >= len(grid[r]):
            continue
        left = join_lines(grid[r][left_col])
        right = join_lines(grid[r][right_col])
        left = re.sub(r"^\(?[a-z]\)\s*", "", left, flags=re.IGNORECASE).strip()
        left = re.sub(r"^[a-z]\.\s*", "", left, flags=re.IGNORECASE).strip()
        if not _pair_is_valid_v1_exact(left, right):
            continue
        pairs.append({"left": left, "right": right})
    seen: set[str] = set()
    out: list[dict] = []
    for p in pairs:
        k = normalize_key(p["left"]) + "=>" + normalize_key(p["right"])
        if k not in seen:
            seen.add(k)
            out.append(p)
    return out


def _iter_elements_recursive_v1_exact(container):
    if hasattr(container, "element"):
        parent_elm = container.element.body
        parent_obj = container
    else:
        parent_elm = container._tc
        parent_obj = container
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield ("p", Paragraph(child, parent_obj))
        elif child.tag.endswith("}tbl"):
            tbl = Table(child, parent_obj)
            yield ("tbl", tbl)
            for row in tbl.rows:
                for cell in row.cells:
                    yield from _iter_elements_recursive_v1_exact(cell)


def _cell_fill_hex_v1_exact(cell) -> str | None:
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = (shd.get(qn("w:fill")) or "").strip().upper()
    return None if fill in ("AUTO", "FFFFFF") else fill or None


def _guess_header_skip_by_row_color(table: Table) -> int:
    if not table.rows:
        return 0
    sig0 = tuple(_cell_fill_hex_v1_exact(c) for c in table.rows[0].cells)
    if not any(sig0):
        return 0
    sample = table.rows[1: min(6, len(table.rows))]
    non_header_like = sum(1 for r in sample if not any(_cell_fill_hex_v1_exact(c) for c in r.cells))
    return 1 if non_header_like >= max(1, len(sample) - 1) else 0


def _pick_term_column_by_fill(table: Table) -> int | None:
    stats: dict[int, int] = {}
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if _cell_fill_hex_v1_exact(cell):
                stats[ci] = stats.get(ci, 0) + 1
    if not stats:
        return None
    best_col = max(stats, key=lambda c: stats[c])
    best = stats[best_col]
    sorted_counts = sorted(stats.values(), reverse=True)
    second = sorted_counts[1] if len(sorted_counts) > 1 else 0
    return best_col if best >= 2 and best >= second + 2 else None


def parse_matching_questions_doc_order_v1_exact(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    recent_paras: list[str] = []
    seen_tables: set[str] = set()

    def choose_stem() -> str | None:
        for t in reversed(recent_paras[-50:]):
            if _looks_like_matching_stem_v1_exact(t):
                return strip_q_prefix(clean_text(t))
        return None

    for kind, el in _iter_elements_recursive_v1_exact(doc):
        if kind == "p":
            t, _ = paragraph_text_and_is_red(el)
            t = clean_text(t)
            if t:
                recent_paras.append(t)
                if len(recent_paras) > 400:
                    recent_paras = recent_paras[-400:]
            continue

        grid = _table_to_grid_v1_exact(el)
        header_skip = _guess_header_skip_by_row_color(el)
        tfp = _table_fingerprint_v1_exact(grid)
        if tfp in seen_tables:
            continue
        seen_tables.add(tfp)
        if _is_instruction_table_v1_exact(grid):
            continue

        term_col = _pick_term_column_by_fill(el)
        if term_col is not None:
            max_cols = max(len(r) for r in grid)
            best_right, best_score = None, 0
            for b in range(max_cols):
                if b == term_col:
                    continue
                sc = _score_columns_v1_exact(grid, term_col, b)
                if sc > best_score:
                    best_score = sc
                    best_right = b
            if best_right is None or best_score < 2:
                continue
            left_col, right_col = term_col, best_right
        else:
            cols = _pick_best_columns_v1_exact(grid)
            if not cols:
                continue
            left_col, right_col = cols

        pairs = _extract_pairs_v1_exact(grid, left_col, right_col, start_row=header_skip)
        if len(pairs) < 2:
            continue

        stem = choose_stem()
        if not stem:
            header = grid[0] if grid else []
            hL = header[left_col][0] if header and left_col < len(header) and header[left_col] else "Left"
            hR = header[right_col][0] if header and right_col < len(header) and header[right_col] else "Right"
            stem = f"Match each '{hL}' to the correct '{hR}'."

        out.append({
            "question": stem,
            "pairs": pairs,
            "kind": "matching",
            "options": [],
            "correct": [],
            "multi": False,
        })
        recent_paras = []
    return out


# ===========================================================================
# V3 matching parser
# ===========================================================================
V3_MATCHING_STEM_RE = re.compile(
    r"\b(complete\s+the\s+table|drag(?:ging)?\s+and\s+drop(?:ping)?|drag\s+and\s+drop|"
    r"match\s+each|match\s+the\s+following|match\s+.*\s+to\s+the\s+correct|select\s+one.*for\s+each)\b",
    re.IGNORECASE,
)
V3_IGNORE_TABLE_HEADERS = re.compile(
    r"^(?:poultry ingredient|definition|style/method of cooking|poultry type or cut|"
    r"essential characteristics|classical chicken dishes|contemporary chicken dishes)\b",
    re.IGNORECASE,
)


def _v3_join_lines(lines: list[str]) -> str:
    parts = [v3_clean_text(x) for x in (lines or []) if v3_clean_text(x)]
    return "; ".join(parts).strip()


def _v3_table_to_grid(tbl: Table) -> list[list[list[str]]]:
    return [
        [[v3_clean_text(p.text) for p in cell.paragraphs if v3_clean_text(p.text)] for cell in row.cells]
        for row in tbl.rows
    ]


def _v3_table_fingerprint(grid) -> str:
    rows = ["|".join(_v3_join_lines(c) for c in r) for r in grid]
    return v3_normalize_key("||".join(rows))


def _v3_header_contains(grid, *needles: str) -> bool:
    if not grid or not grid[0]:
        return False
    header = " | ".join(_v3_join_lines(c) for c in grid[0]).lower()
    return all(n.lower() in header for n in needles)


def _v3_is_table_forced_essay(grid) -> bool:
    return (
        _v3_header_contains(grid, "poultry ingredient", "definition", "style")
        or _v3_header_contains(grid, "poultry type", "essential")
        or _v3_header_contains(grid, "classical chicken dishes", "contemporary chicken dishes")
    )


def _v3_looks_like_matching_stem(t: str) -> bool:
    t2 = v3_strip_q_prefix(v3_clean_text(t))
    if not t2 or t2.lower().startswith(("for learners", "for assessors", "for students")):
        return False
    return bool(V3_MATCHING_STEM_RE.search(t2))


def _v3_score_columns(grid, a: int, b: int) -> int:
    sc = 0
    for r in grid[1:]:
        if a >= len(r) or b >= len(r):
            continue
        if _v3_join_lines(r[a]) and _v3_join_lines(r[b]):
            sc += 1
    return sc


def _v3_pick_best_columns(grid):
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    best = None
    best_sc = 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = _v3_score_columns(grid, a, b)
            if sc > best_sc:
                best_sc = sc
                best = (a, b)
    return best if best_sc >= 2 else None


def _v3_extract_pairs(grid, left_col: int, right_col: int, start_row: int = 1) -> list[dict]:
    pairs = []
    for r in grid[start_row:]:
        if left_col >= len(r) or right_col >= len(r):
            continue
        left = _v3_join_lines(r[left_col])
        right = _v3_join_lines(r[right_col])
        if left and right:
            pairs.append({"left": left, "right": right})
    return pairs


def _v3_find_item_index(items: list[dict], needle: str) -> int | None:
    n = v3_normalize_key(needle)
    if not n:
        return None
    for i, it in enumerate(items):
        t = v3_normalize_key(it.get("text", ""))
        if t and (t == n or n in t):
            return i
    return None


def _v3_is_instructions_matching(pairs: list[dict], stem: str) -> bool:
    s = v3_normalize_key(stem or "")
    if "instructions" in s and ("for students" in s or "for learners" in s or "for assessors" in s):
        return True
    left_keys = {v3_normalize_key((p.get("left") or "")) for p in (pairs or [])}
    common = {
        "range and conditions", "decision-making rules", "decision making rules",
        "pre-approved reasonable adjustments", "pre approved reasonable adjustments",
        "rubric", "instructions",
    }
    if sum(1 for k in common if k in left_keys) >= 3:
        return True
    right_blob = v3_normalize_key("; ".join(p.get("right") or "" for p in (pairs or [])))
    return any(
        phrase in right_blob
        for phrase in [
            "students must work through this assessment independently",
            "false declarations may lead to withdrawal",
            "feedback comments must be provided",
        ]
    )


def v3_parse_matching_questions_doc_order(docx_path: str, items: list[dict] | None = None) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    seen: set[str] = set()
    seq = 0

    for el in v3_iter_all_tables(doc):
        seq += 1
        grid = _v3_table_to_grid(el)
        if _v3_is_table_forced_essay(grid):
            continue
        fp = _v3_table_fingerprint(grid)
        if fp in seen:
            continue
        seen.add(fp)

        cols = _v3_pick_best_columns(grid)
        if not cols:
            continue
        left_col, right_col = cols
        pairs = _v3_extract_pairs(grid, left_col, right_col, start_row=1)
        if len(pairs) < 2:
            continue

        header = grid[0] if grid else []
        hL = _v3_join_lines(header[left_col]) if header and left_col < len(header) else "Left"
        hR = _v3_join_lines(header[right_col]) if header and right_col < len(header) else "Right"
        stem = f"Match each '{hL}' to the correct '{hR}'."

        if _v3_is_instructions_matching(pairs, stem):
            continue

        order = seq
        if items:
            for cand in (hL, hR, pairs[0].get("left", ""), pairs[0].get("right", "")):
                idx = _v3_find_item_index(items, cand)
                if idx is not None:
                    order = idx
                    break

        out.append({
            "question": stem,
            "pairs": pairs,
            "kind": "matching",
            "options": [],
            "correct": [],
            "multi": False,
            "_order": order,
            "qnum": None,
        })
    return out


# ===========================================================================
# V3 table-to-essay helpers (forced essay tables)
# ===========================================================================
def v3_parse_table_defined_terms_as_essays(docx_path: str, items: list[dict]) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    seen_terms: set[str] = set()
    for el in v3_iter_all_tables(doc):
        grid = _v3_table_to_grid(el)
        if not _v3_header_contains(grid, "poultry ingredient", "definition", "style"):
            continue
        for r in grid[1:]:
            if not r:
                continue
            term = v3_strip_q_prefix(_v3_join_lines(r[0] if r else []))
            term = v3_clean_text(re.sub(r"^\s*(?:q\s*)?\d+\s*[\.\)]\s*", "", term, flags=re.IGNORECASE))
            if not term or term.lower() in {"poultry ingredient", "definition", "style/method of cooking"}:
                continue
            k = v3_normalize_key(term)
            if not k or k in seen_terms:
                continue
            seen_terms.add(k)
            order = _v3_find_item_index(items, term) or _v3_find_item_index(items, f"Define {term}") or 10**9
            out.append({
                "question": f"Define: {term}. Provide one style/method of cooking.",
                "options": [], "correct": [], "multi": False, "kind": "essay",
                "_order": order, "qnum": None,
            })
    return out


def v3_parse_table_characteristics_as_essays(docx_path: str, items: list[dict]) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    seen_terms: set[str] = set()
    for el in v3_iter_all_tables(doc):
        grid = _v3_table_to_grid(el)
        if not _v3_header_contains(grid, "poultry type", "essential"):
            continue
        for r in grid[1:]:
            if not r:
                continue
            term = v3_strip_q_prefix(_v3_join_lines(r[0] if r else []))
            term = v3_clean_text(re.sub(r"^\s*(?:q\s*)?\d+\s*[\.\)]\s*", "", term, flags=re.IGNORECASE))
            if not term or term.lower() in {"poultry type or cut", "essential characteristics"}:
                continue
            k = v3_normalize_key(term)
            if not k or k in seen_terms:
                continue
            seen_terms.add(k)
            order = _v3_find_item_index(items, term) or _v3_find_item_index(items, f"Describe {term}") or 10**9
            out.append({
                "question": f"Describe the essential characteristics of: {term}.",
                "options": [], "correct": [], "multi": False, "kind": "essay",
                "_order": order, "qnum": None,
            })
    return out


def v3_collect_ignore_texts_from_forced_tables(docx_path: str) -> set[str]:
    doc = Document(docx_path)
    ignore: set[str] = set()

    def add_lines(lines: list[str]):
        for ln in lines or []:
            t = v3_clean_text(ln)
            if not t or (len(t) < 12 and " " not in t) or re.fullmatch(r"\d+", t):
                continue
            ignore.add(t)

    for tbl in v3_iter_all_tables(doc):
        grid = _v3_table_to_grid(tbl)
        if _v3_header_contains(grid, "poultry ingredient", "definition", "style"):
            for row in grid[1:]:
                for cell in row[1:]:
                    add_lines(cell)
        elif _v3_header_contains(grid, "poultry type", "essential"):
            for row in grid[1:]:
                for cell in row[1:]:
                    add_lines(cell)
        elif _v3_header_contains(grid, "classical chicken dishes", "contemporary chicken dishes"):
            for row in grid[1:]:
                for cell in row:
                    add_lines(cell)
    return ignore