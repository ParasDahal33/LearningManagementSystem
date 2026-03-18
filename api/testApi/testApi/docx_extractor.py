"""
parsers/docx_extractor.py
Low-level DOCX extraction: paragraph text, red-run detection,
textbox / DrawingML text, table iteration, and version-specific extractors
(v1, v2, v3).
"""

from __future__ import annotations

import re
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .util import (
    clean_text,
    is_red_hex,
    v2_clean_text,
    v3_clean_text,
)


# ---------------------------------------------------------------------------
# Shared: run / paragraph helpers
# ---------------------------------------------------------------------------
def is_red_run(run) -> bool:
    color = run.font.color
    if not color:
        return False
    rgb = color.rgb
    if rgb is None:
        return False
    r, g, b = rgb[0], rgb[1], rgb[2]
    return r >= 200 and g <= 80 and b <= 80


def paragraph_text_and_is_red(paragraph: Paragraph) -> tuple[str, bool]:
    text = "".join(run.text for run in paragraph.runs).strip()
    any_red = any(is_red_run(run) and run.text.strip() for run in paragraph.runs)
    return text, any_red


def local(tag: str) -> str:
    return (tag or "").rsplit("}", 1)[-1]


# ---------------------------------------------------------------------------
# Shared: textbox and DrawingML helpers (v1 version)
# ---------------------------------------------------------------------------
def txbx_paragraph_text_and_is_red(p_elm: CT_P) -> tuple[str, bool]:
    parts: list[str] = []
    any_red = False
    for r in p_elm.iter():
        if not str(getattr(r, "tag", "")).endswith("}r"):
            continue
        run_texts = [
            t_node.text
            for t_node in r.iter()
            if str(getattr(t_node, "tag", "")).endswith("}t") and getattr(t_node, "text", None)
        ]
        t = "".join(run_texts)
        if t:
            parts.append(t)
        if any_red or not t.strip():
            continue
        rpr = next(
            (child for child in list(r) if str(getattr(child, "tag", "")).endswith("}rPr")),
            None,
        )
        if rpr is None:
            continue
        color = next(
            (c for c in rpr.iter() if str(getattr(c, "tag", "")).endswith("}color")),
            None,
        )
        if color is None:
            continue
        val = next(
            (v for k, v in getattr(color, "attrib", {}).items() if str(k).endswith("}val") or str(k) == "val"),
            None,
        )
        if val and is_red_hex(val):
            any_red = True
    return clean_text("".join(parts)), any_red


def textbox_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []
    out: list[tuple[str, bool]] = []
    for el in p_elm.iter():
        if not str(getattr(el, "tag", "")).endswith("}txbxContent"):
            continue
        for p2 in el.iter():
            if str(getattr(p2, "tag", "")).endswith("}p"):
                t, red = txbx_paragraph_text_and_is_red(p2)
                if t:
                    out.append((t, red))
    return out


def drawingml_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []
    lines: list[tuple[str, bool]] = []
    seen: set = set()
    for el in p_elm.iter():
        tag = str(getattr(el, "tag", ""))
        if not tag.endswith("}p"):
            continue
        ns = tag.split("}")[0].lstrip("{")
        if "drawing" not in ns.lower():
            continue
        parts: list[str] = []
        any_red = False
        for node in el.iter():
            ntag = str(getattr(node, "tag", ""))
            if ntag.endswith("}t") and getattr(node, "text", None):
                parts.append(node.text)
            if not any_red and ntag.endswith("}srgbClr"):
                val = next(
                    (v for k, v in getattr(node, "attrib", {}).items() if str(k).endswith("}val") or str(k) == "val"),
                    None,
                )
                if val and is_red_hex(val):
                    any_red = True
        text = clean_text("".join(parts))
        if not text:
            continue
        key = (text, any_red)
        if key in seen:
            continue
        seen.add(key)
        lines.append((text, any_red))
    return lines


# ---------------------------------------------------------------------------
# Shared: block iteration
# ---------------------------------------------------------------------------
def iter_block_items(doc: Document):
    """Yield Paragraph or Table objects from the document body in order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


# ---------------------------------------------------------------------------
# v1 extractors
# ---------------------------------------------------------------------------
def extract_items_with_red(docx_path: str) -> list[dict]:
    """Full extraction including textbox and DrawingML content (v1)."""
    doc = Document(docx_path)
    items: list[dict] = []

    def push_text(t: str, is_red: bool):
        t = clean_text(t)
        if t:
            items.append({"text": t, "is_red": is_red})

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t, red = paragraph_text_and_is_red(block)
            push_text(t, red)
            for t2, red2 in textbox_texts_in_paragraph(block):
                push_text(t2, red2)
            for t2, red2 in drawingml_texts_in_paragraph(block):
                push_text(t2, red2)
        else:
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t, red = paragraph_text_and_is_red(p)
                        push_text(t, red)
                        for t2, red2 in textbox_texts_in_paragraph(p):
                            push_text(t2, red2)
                        for t2, red2 in drawingml_texts_in_paragraph(p):
                            push_text(t2, red2)
    return items


def extract_items_with_red_v1(docx_path: str) -> list[dict]:
    """Paragraph-only extraction (stable, used for description building)."""
    doc = Document(docx_path)
    items: list[dict] = []

    def push_text(t: str, is_red: bool):
        t = clean_text(t)
        if t:
            items.append({"text": t, "is_red": bool(is_red)})

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t, red = paragraph_text_and_is_red(block)
            push_text(t, red)
        else:
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t, red = paragraph_text_and_is_red(p)
                        push_text(t, red)
    return items


# ---------------------------------------------------------------------------
# v2 extractor
# ---------------------------------------------------------------------------
def _v2_txbx_paragraph_text_and_is_red(p_elm: CT_P) -> tuple[str, bool]:
    parts: list[str] = []
    any_red = False
    for r in p_elm.iter():
        if not str(getattr(r, "tag", "")).endswith("}r"):
            continue
        run_texts = [
            t_node.text
            for t_node in r.iter()
            if str(getattr(t_node, "tag", "")).endswith("}t") and getattr(t_node, "text", None)
        ]
        t = "".join(run_texts)
        if t:
            parts.append(t)
        if any_red or not t.strip():
            continue
        rpr = next(
            (child for child in list(r) if str(getattr(child, "tag", "")).endswith("}rPr")),
            None,
        )
        if rpr is None:
            continue
        color = next(
            (c for c in rpr.iter() if str(getattr(c, "tag", "")).endswith("}color")),
            None,
        )
        if color is None:
            continue
        val = next(
            (v for k, v in getattr(color, "attrib", {}).items() if str(k).endswith("}val") or str(k) == "val"),
            None,
        )
        if val and is_red_hex(val):
            any_red = True
    return v2_clean_text("".join(parts)), any_red


def _v2_textbox_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []
    out: list[tuple[str, bool]] = []
    for el in p_elm.iter():
        if not str(getattr(el, "tag", "")).endswith("}txbxContent"):
            continue
        for p2 in el.iter():
            if str(getattr(p2, "tag", "")).endswith("}p"):
                t, red = _v2_txbx_paragraph_text_and_is_red(p2)
                if t:
                    out.append((t, red))
    return out


def _v2_drawingml_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []
    lines: list[tuple[str, bool]] = []
    seen: set = set()
    for el in p_elm.iter():
        tag = str(getattr(el, "tag", ""))
        if not tag.endswith("}p"):
            continue
        ns = tag.split("}")[0].lstrip("{")
        if "drawing" not in ns.lower():
            continue
        parts: list[str] = []
        any_red = False
        for node in el.iter():
            ntag = str(getattr(node, "tag", ""))
            if ntag.endswith("}t") and getattr(node, "text", None):
                parts.append(node.text)
            if not any_red and ntag.endswith("}srgbClr"):
                val = next(
                    (v for k, v in getattr(node, "attrib", {}).items() if str(k).endswith("}val") or str(k) == "val"),
                    None,
                )
                if val and is_red_hex(val):
                    any_red = True
        text = v2_clean_text("".join(parts))
        if not text:
            continue
        key = (text, any_red)
        if key in seen:
            continue
        seen.add(key)
        lines.append((text, any_red))
    return lines


def v2_extract_items_with_red(docx_path: str) -> list[dict]:
    """
    v2 extractor: prefers python-docx; falls back to raw XML.
    Picks whichever extraction scores higher on question-verb heuristics.
    """

    def extract_with_python_docx(path: str) -> list[dict]:
        doc = Document(path)
        items: list[dict] = []
        last = None

        def push(t: str, red: bool):
            nonlocal last
            t = v2_clean_text(t)
            if not t:
                return
            key = (t, bool(red))
            if last == key:
                return
            items.append({"text": t, "is_red": bool(red)})
            last = key

        def push_paragraph(p: Paragraph):
            t, red = paragraph_text_and_is_red(p)
            push(t, red)
            for t2, red2 in _v2_textbox_texts_in_paragraph(p):
                push(t2, red2)
            for t2, red2 in _v2_drawingml_texts_in_paragraph(p):
                push(t2, red2)

        def push_table(tbl: Table):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        push_paragraph(p)
                    for sub in getattr(cell, "tables", []) or []:
                        push_table(sub)

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                push_paragraph(block)
            else:
                push_table(block)
        return items

    def extract_with_xml_fallback(path: str) -> list[dict]:
        try:
            with zipfile.ZipFile(path) as z:
                raw = z.read("word/document.xml")
        except Exception:
            return []
        try:
            root = ET.fromstring(raw)
        except Exception:
            return []

        def run_is_red(run_el) -> bool:
            for el in run_el.iter():
                lname = local(getattr(el, "tag", ""))
                if lname in ("color", "srgbClr"):
                    for k, v in getattr(el, "attrib", {}).items():
                        if (str(k).endswith("}val") or str(k) == "val") and v and is_red_hex(v):
                            return True
            return False

        items: list[dict] = []
        last = None

        def push(parts: list[str], red: bool):
            nonlocal last
            t = v2_clean_text("".join(parts))
            if not t:
                return
            key = (t, bool(red))
            if last == key:
                return
            items.append({"text": t, "is_red": bool(red)})
            last = key

        def walk(node, in_p: bool = False, buf: list[str] | None = None, red_any: list[bool] | None = None):
            lname = local(str(getattr(node, "tag", "")))
            if lname == "p":
                b: list[str] = []
                r = [False]
                for ch in list(node):
                    walk(ch, in_p=True, buf=b, red_any=r)
                push(b, r[0])
                return
            if lname == "r":
                is_red = run_is_red(node)
                for ch in list(node):
                    walk(ch, in_p=in_p, buf=buf, red_any=red_any)
                if is_red and red_any is not None:
                    red_any[0] = True
                return
            if lname == "t":
                if in_p and buf is not None and getattr(node, "text", None):
                    buf.append(node.text)
                return
            for ch in list(node):
                walk(ch, in_p=in_p, buf=buf, red_any=red_any)

        for ch in list(root):
            walk(ch)
        return items

    def score(items: list[dict]) -> int:
        if not items:
            return 0
        qverb = re.compile(
            r"^(q\s*\d+\s*[\.)]\s*)?(list|describe|explain|outline|state|name|provide|define|identify|select|choose|pick|match|complete)\b",
            re.IGNORECASE,
        )
        sc = 0
        for it in items:
            t = v2_clean_text(it.get("text", ""))
            if not t:
                continue
            if qverb.match(t) or "Which of the following" in t:
                sc += 3
            if t.endswith("?"):
                sc += 1
        return sc + min(80, len(items) // 8)

    items_docx = extract_with_python_docx(docx_path)
    items_xml = extract_with_xml_fallback(docx_path)
    return items_xml if score(items_xml) > score(items_docx) else items_docx


# ---------------------------------------------------------------------------
# v3 extractor
# ---------------------------------------------------------------------------
def _v3_txbx_paragraph_text_and_is_red(p_elm: CT_P) -> tuple[str, bool]:
    parts: list[str] = []
    any_red = False
    for r in p_elm.iter():
        if not str(getattr(r, "tag", "")).endswith("}r"):
            continue
        run_texts = [
            t_node.text
            for t_node in r.iter()
            if str(getattr(t_node, "tag", "")).endswith("}t") and getattr(t_node, "text", None)
        ]
        t = "".join(run_texts)
        if t:
            parts.append(t)
        if any_red or not t.strip():
            continue
        rpr = next(
            (child for child in list(r) if str(getattr(child, "tag", "")).endswith("}rPr")),
            None,
        )
        if rpr is None:
            continue
        color = next(
            (c for c in rpr.iter() if str(getattr(c, "tag", "")).endswith("}color")),
            None,
        )
        if color is None:
            continue
        val = next(
            (v for k, v in getattr(color, "attrib", {}).items() if str(k).endswith("}val") or str(k) == "val"),
            None,
        )
        if val and is_red_hex(val):
            any_red = True
    return v3_clean_text("".join(parts)), any_red


def _v3_textbox_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []
    out: list[tuple[str, bool]] = []
    for el in p_elm.iter():
        if not str(getattr(el, "tag", "")).endswith("}txbxContent"):
            continue
        for p2 in el.iter():
            if str(getattr(p2, "tag", "")).endswith("}p"):
                t, red = _v3_txbx_paragraph_text_and_is_red(p2)
                if t:
                    out.append((t, red))
    return out


def _v3_drawingml_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []
    lines: list[tuple[str, bool]] = []
    seen: set = set()
    for el in p_elm.iter():
        tag = str(getattr(el, "tag", ""))
        if not tag.endswith("}p"):
            continue
        ns = tag.split("}")[0].lstrip("{")
        if "drawing" not in ns.lower():
            continue
        parts: list[str] = []
        any_red = False
        for node in el.iter():
            ntag = str(getattr(node, "tag", ""))
            if ntag.endswith("}t") and getattr(node, "text", None):
                parts.append(node.text)
            if not any_red and ntag.endswith("}srgbClr"):
                val = next(
                    (v for k, v in getattr(node, "attrib", {}).items() if str(k).endswith("}val") or str(k) == "val"),
                    None,
                )
                if val and is_red_hex(val):
                    any_red = True
        text = v3_clean_text("".join(parts))
        if not text:
            continue
        key = (text, any_red)
        if key in seen:
            continue
        seen.add(key)
        lines.append((text, any_red))
    return lines


def v3_extract_items_with_red(docx_path: str, include_tables: bool = True) -> list[dict]:
    """v3 extractor: python-docx only, tags items with src='body'|'table'."""

    def extract_with_python_docx(path: str) -> list[dict]:
        doc = Document(path)
        items: list[dict] = []
        last = None

        def push(t: str, red: bool, src: str):
            nonlocal last
            t = v3_clean_text(t)
            if not t:
                return
            key = (t, bool(red), src)
            if last == key:
                return
            items.append({"text": t, "is_red": bool(red), "src": src})
            last = key

        def push_paragraph(p: Paragraph, src: str):
            t, red = paragraph_text_and_is_red(p)
            push(t, red, src)
            for t2, red2 in _v3_textbox_texts_in_paragraph(p):
                push(t2, red2, src)
            for t2, red2 in _v3_drawingml_texts_in_paragraph(p):
                push(t2, red2, src)

        def push_table(tbl: Table):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        push_paragraph(p, "table")
                    for sub in getattr(cell, "tables", []) or []:
                        push_table(sub)

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                push_paragraph(block, "body")
            elif include_tables:
                push_table(block)
        return items

    return extract_with_python_docx(docx_path)


# ---------------------------------------------------------------------------
# v3: table helpers used by matching and essay parsers
# ---------------------------------------------------------------------------
def v3_iter_tables_recursive(tbl: Table):
    yield tbl
    for row in tbl.rows:
        for cell in row.cells:
            for sub in getattr(cell, "tables", []) or []:
                yield from v3_iter_tables_recursive(sub)


def v3_iter_all_tables(doc: Document):
    for el in iter_block_items(doc):
        if isinstance(el, Table):
            yield from v3_iter_tables_recursive(el)


# ---------------------------------------------------------------------------
# v1: description builder
# ---------------------------------------------------------------------------
QUESTION_CMD_INNER_RE_V1 = re.compile(
    r"\b(Which\s+of\s+the\s+following\b|"
    r"(Identify|Select)\s+(one|two|three|four|five|six|seven|eight|nine|ten|\d+)\b)",
    re.IGNORECASE,
)
COMMAND_QUESTION_RE_V1 = re.compile(
    r"^(Illustrate|Critically\s+(?:assess|analyse|analyze|evaluate)|"
    r"Evaluate|Determine|Articulate|Prescribe|Analyse|Analyze|Review|Recommend)\b.+",
    re.IGNORECASE,
)


def build_description_v1(items: list[dict]) -> str:
    """Build the quiz description HTML block from document preamble."""
    from core.utils import strip_q_prefix

    collecting = False
    lines = []
    for it in items:
        t = clean_text(it.get("text", ""))
        if not t:
            continue
        if re.search(r"\bFor learners\b", t, re.IGNORECASE):
            collecting = True
        if collecting and (
            QUESTION_CMD_INNER_RE_V1.search(t)
            or COMMAND_QUESTION_RE_V1.match(strip_q_prefix(t))
            or re.search(
                r"\bdragging\s+and\s+dropping\b|\bdrag\s+and\s+drop\b|\bComplete\s+the\s+table\b",
                t,
                re.IGNORECASE,
            )
        ):
            break
        if collecting:
            lines.append(t)
    if not lines:
        return ""

    html_parts: list[str] = []
    in_list = False
    for ln in lines:
        ln = ln.strip()
        if "•" in ln and not ln.strip().startswith("•"):
            before, *bullets = [p.strip() for p in ln.split("•") if p.strip()]
            if before:
                if in_list:
                    html_parts.append("</ul>")
                    in_list = False
                html_parts.append(f"<p>{before}</p>")
            if bullets:
                if not in_list:
                    html_parts.append("<ul>")
                    in_list = True
                for b in bullets:
                    html_parts.append(f"<li>{b}</li>")
            continue
        if ln.startswith("•"):
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{ln.lstrip('•').strip()}</li>")
            continue
        if in_list:
            html_parts.append("</ul>")
            in_list = False
        html_parts.append(f"<p>{ln}</p>")
    if in_list:
        html_parts.append("</ul>")
    return "\n".join(html_parts)