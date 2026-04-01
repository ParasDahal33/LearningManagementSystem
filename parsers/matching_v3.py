"""
parsers/matching_v3.py
V3 matching parsers and forced table-to-essay converters.
"""

from __future__ import annotations

import re

from docx import Document
from docx.table import Table

from core.utils import (
    v3_clean_text,
    v3_normalize_key,
    v3_strip_q_prefix,
)
from .docx_extractor import v3_iter_all_tables


# ===========================================================================
# V3 matching parser
# ===========================================================================
V3_MATCHING_STEM_RE = re.compile(
    r"\b(complete\s+the\s+table|drag(?:ging)?\s+and\s+drop(?:ping)?|drag\s+and\s+drop|"
    r"match\s+each|match\s+the\s+following|match\s+.*\s+to\s+the\s+correct|select\s+one.*for\s+each)\b",
    re.IGNORECASE,
)
V3_IGNORE_TABLE_HEADERS = re.compile(
    r"^(?:poultry ingredient|definition|style/method of cooking|poultry type or cut|"
    r"essential characteristics|classical chicken dishes|contemporary chicken dishes)\b",
    re.IGNORECASE,
)


def _v3_join_lines(lines: list[str]) -> str:
    parts = [v3_clean_text(x) for x in (lines or []) if v3_clean_text(x)]
    return "; ".join(parts).strip()


def _v3_table_to_grid(tbl: Table) -> list[list[list[str]]]:
    return [
        [[v3_clean_text(p.text) for p in cell.paragraphs if v3_clean_text(p.text)] for cell in row.cells]
        for row in tbl.rows
    ]


def _v3_table_fingerprint(grid) -> str:
    rows = ["|".join(_v3_join_lines(c) for c in r) for r in grid]
    return v3_normalize_key("||".join(rows))


def _v3_header_contains(grid, *needles: str) -> bool:
    if not grid or not grid[0]:
        return False
    header = " | ".join(_v3_join_lines(c) for c in grid[0]).lower()
    return all(n.lower() in header for n in needles)


def _v3_is_table_forced_essay(grid) -> bool:
    return (
        _v3_header_contains(grid, "poultry ingredient", "definition", "style")
        or _v3_header_contains(grid, "poultry type", "essential")
        or _v3_header_contains(grid, "classical chicken dishes", "contemporary chicken dishes")
    )


def _v3_looks_like_matching_stem(t: str) -> bool:
    t2 = v3_strip_q_prefix(v3_clean_text(t))
    if not t2 or t2.lower().startswith(("for learners", "for assessors", "for students")):
        return False
    return bool(V3_MATCHING_STEM_RE.search(t2))


def _v3_score_columns(grid, a: int, b: int) -> int:
    sc = 0
    for r in grid[1:]:
        if a >= len(r) or b >= len(r):
            continue
        if _v3_join_lines(r[a]) and _v3_join_lines(r[b]):
            sc += 1
    return sc


def _v3_pick_best_columns(grid):
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    best = None
    best_sc = 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = _v3_score_columns(grid, a, b)
            if sc > best_sc:
                best_sc = sc
                best = (a, b)
    return best if best_sc >= 2 else None


def _v3_extract_pairs(grid, left_col: int, right_col: int, start_row: int = 1) -> list[dict]:
    pairs = []
    for r in grid[start_row:]:
        if left_col >= len(r) or right_col >= len(r):
            continue
        left = _v3_join_lines(r[left_col])
        right = _v3_join_lines(r[right_col])
        if left and right:
            pairs.append({"left": left, "right": right})
    return pairs


def _v3_find_item_index(items: list[dict], needle: str) -> int | None:
    n = v3_normalize_key(needle)
    if not n:
        return None
    for i, it in enumerate(items):
        t = v3_normalize_key(it.get("text", ""))
        if t and (t == n or n in t):
            return i
    return None


def _v3_is_instructions_matching(pairs: list[dict], stem: str) -> bool:
    s = v3_normalize_key(stem or "")
    if "instructions" in s and ("for students" in s or "for learners" in s or "for assessors" in s):
        return True
    left_keys = {v3_normalize_key((p.get("left") or "")) for p in (pairs or [])}
    common = {
        "range and conditions", "decision-making rules", "decision making rules",
        "pre-approved reasonable adjustments", "pre approved reasonable adjustments",
        "rubric", "instructions",
    }
    if sum(1 for k in common if k in left_keys) >= 3:
        return True
    right_blob = v3_normalize_key("; ".join(p.get("right") or "" for p in (pairs or [])))
    return any(
        phrase in right_blob
        for phrase in [
            "students must work through this assessment independently",
            "false declarations may lead to withdrawal",
            "feedback comments must be provided",
        ]
    )


def v3_parse_matching_questions_doc_order(docx_path: str, items: list[dict] | None = None) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    seen: set[str] = set()
    seq = 0

    for el in v3_iter_all_tables(doc):
        seq += 1
        grid = _v3_table_to_grid(el)
        if _v3_is_table_forced_essay(grid):
            continue
        fp = _v3_table_fingerprint(grid)
        if fp in seen:
            continue
        seen.add(fp)

        cols = _v3_pick_best_columns(grid)
        if not cols:
            continue
        left_col, right_col = cols
        pairs = _v3_extract_pairs(grid, left_col, right_col, start_row=1)
        if len(pairs) < 2:
            continue

        header = grid[0] if grid else []
        hL = _v3_join_lines(header[left_col]) if header and left_col < len(header) else "Left"
        hR = _v3_join_lines(header[right_col]) if header and right_col < len(header) else "Right"
        stem = f"Match each '{hL}' to the correct '{hR}'."

        if _v3_is_instructions_matching(pairs, stem):
            continue

        order = seq
        if items:
            for cand in (hL, hR, pairs[0].get("left", ""), pairs[0].get("right", "")):
                idx = _v3_find_item_index(items, cand)
                if idx is not None:
                    order = idx
                    break

        out.append({
            "question": stem,
            "pairs": pairs,
            "kind": "matching",
            "options": [],
            "correct": [],
            "multi": False,
            "_order": order,
            "qnum": None,
        })
    return out


# ===========================================================================
# V3 table-to-essay helpers (forced essay tables)
# ===========================================================================
def v3_parse_table_defined_terms_as_essays(docx_path: str, items: list[dict]) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    seen_terms: set[str] = set()
    for el in v3_iter_all_tables(doc):
        grid = _v3_table_to_grid(el)
        if not _v3_header_contains(grid, "poultry ingredient", "definition", "style"):
            continue
        for r in grid[1:]:
            if not r:
                continue
            term = v3_strip_q_prefix(_v3_join_lines(r[0] if r else []))
            term = v3_clean_text(re.sub(r"^\s*(?:q\s*)?\d+\s*[\.\)]\s*", "", term, flags=re.IGNORECASE))
            if not term or term.lower() in {"poultry ingredient", "definition", "style/method of cooking"}:
                continue
            k = v3_normalize_key(term)
            if not k or k in seen_terms:
                continue
            seen_terms.add(k)
            order = _v3_find_item_index(items, term) or _v3_find_item_index(items, f"Define {term}") or 10**9
            out.append({
                "question": f"Define: {term}. Provide one style/method of cooking.",
                "options": [], "correct": [], "multi": False, "kind": "essay",
                "_order": order, "qnum": None,
            })
    return out


def v3_parse_table_characteristics_as_essays(docx_path: str, items: list[dict]) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    seen_terms: set[str] = set()
    for el in v3_iter_all_tables(doc):
        grid = _v3_table_to_grid(el)
        if not _v3_header_contains(grid, "poultry type", "essential"):
            continue
        for r in grid[1:]:
            if not r:
                continue
            term = v3_strip_q_prefix(_v3_join_lines(r[0] if r else []))
            term = v3_clean_text(re.sub(r"^\s*(?:q\s*)?\d+\s*[\.\)]\s*", "", term, flags=re.IGNORECASE))
            if not term or term.lower() in {"poultry type or cut", "essential characteristics"}:
                continue
            k = v3_normalize_key(term)
            if not k or k in seen_terms:
                continue
            seen_terms.add(k)
            order = _v3_find_item_index(items, term) or _v3_find_item_index(items, f"Describe {term}") or 10**9
            out.append({
                "question": f"Describe the essential characteristics of: {term}.",
                "options": [], "correct": [], "multi": False, "kind": "essay",
                "_order": order, "qnum": None,
            })
    return out


def v3_collect_ignore_texts_from_forced_tables(docx_path: str) -> set[str]:
    doc = Document(docx_path)
    ignore: set[str] = set()

    def add_lines(lines: list[str]):
        for ln in lines or []:
            t = v3_clean_text(ln)
            if not t or (len(t) < 12 and " " not in t) or re.fullmatch(r"\d+", t):
                continue
            ignore.add(t)

    for tbl in v3_iter_all_tables(doc):
        grid = _v3_table_to_grid(tbl)
        if _v3_header_contains(grid, "poultry ingredient", "definition", "style"):
            for row in grid[1:]:
                for cell in row[1:]:
                    add_lines(cell)
        elif _v3_header_contains(grid, "poultry type", "essential"):
            for row in grid[1:]:
                for cell in row[1:]:
                    add_lines(cell)
        elif _v3_header_contains(grid, "classical chicken dishes", "contemporary chicken dishes"):
            for row in grid[1:]:
                for cell in row:
                    add_lines(cell)
    return ignore