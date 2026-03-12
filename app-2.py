
from __future__ import annotations

import contextlib
import hashlib
import io
import math
import os
import re
import tempfile
import tomllib
import zipfile
from dataclasses import dataclass
from datetime import date, datetime, time
from typing import Any
from xml.etree import ElementTree as ET

import pytz
import requests
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph



TZ_NAME = "Australia/Sydney"
tz = pytz.timezone(TZ_NAME)



def safe_load_secrets_toml() -> dict[str, str]:
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        secrets_path = os.path.join(base_dir, ".streamlit", "secrets.toml")
        if not os.path.exists(secrets_path):
            return {}
        with open(secrets_path, "rb") as f:
            data = tomllib.load(f)
        out: dict[str, str] = {}
        for k in ["CANVAS_BASE_URL", "CANVAS_TOKEN", "OPENAI_API_KEY", "OPENAI_MODEL", "OPENAI_BASE_URL"]:
            v = data.get(k)
            if isinstance(v, str):
                out[k] = v
        return out
    except Exception:
        return {}


_LOCAL_SECRETS = safe_load_secrets_toml()



def ss_init(key: str, value: Any) -> None:
    if key not in st.session_state:
        st.session_state[key] = value


ss_init("logged_in", False)
ss_init("me", None)
ss_init("canvas_token", os.getenv("CANVAS_TOKEN", "") or _LOCAL_SECRETS.get("CANVAS_TOKEN", ""))
ss_init(
    "canvas_base_url",
    os.getenv("CANVAS_BASE_URL", "") or _LOCAL_SECRETS.get("CANVAS_BASE_URL", "https://learningvault.test.instructure.com/api/v1"),
)
ss_init("courses_cache", None)
ss_init("selected_course_id", None)

ss_init("docx_filename", None)
ss_init("description_html", "")
ss_init("questions", [])
ss_init("parsed_ok", False)
ss_init("parse_run_id", 0)
ss_init("last_parser_mode", None)
ss_init("questions_page_size", 10)
ss_init("questions_page", 1)

ss_init(
    "details",
    {
        "shuffle_answers": True,
        "time_limit": 0,
        "allow_multiple_attempts": False,
        "allowed_attempts": 2,
        "scoring_policy": "keep_highest",
        "one_question_at_a_time": False,
        "show_correct_answers": False,
        "access_code_enabled": False,
        "access_code": "",
        "due_at": "",
        "unlock_at": "",
        "lock_at": "",
    },
)

ss_init("openai_api_key", os.getenv("OPENAI_API_KEY", "") or _LOCAL_SECRETS.get("OPENAI_API_KEY", ""))
ss_init("openai_model", os.getenv("OPENAI_MODEL", "") or _LOCAL_SECRETS.get("OPENAI_MODEL", "gpt-4.1-mini"))
ss_init("openai_base_url", os.getenv("OPENAI_BASE_URL", "") or _LOCAL_SECRETS.get("OPENAI_BASE_URL", "https://api.openai.com"))


def combine_date_time(d: date | None, t: time | None) -> str | None:
    if not d or not t:
        return None
    dt = tz.localize(datetime.combine(d, t))
    return dt.isoformat()


def clean_text(t: str) -> str:
    return re.sub(r"\s+", " ", (t or "")).strip()


def normalize_key(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    return s


Q_PREFIX_RE_V1 = re.compile(r"^[^A-Za-z0-9]*(?:lo\s*)?Q\s*\d+\s*[\.\)]\s*", re.IGNORECASE)


def strip_q_prefix(line: str) -> str:
    return Q_PREFIX_RE_V1.sub("", (line or "").strip()).strip()


def question_fingerprint(q: dict) -> str:
    qt = normalize_key(q.get("question", ""))
    kind = normalize_key(q.get("kind", ""))
    opts = [normalize_key(x) for x in (q.get("options") or [])]
    pairs = q.get("pairs") or []
    pairs_blob = "||".join([normalize_key(p.get("left", "")) + "=>" + normalize_key(p.get("right", "")) for p in pairs])
    blob = kind + "||" + qt + "||" + "||".join(opts) + "||" + pairs_blob
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()


def dedupe_questions(questions: list[dict]) -> list[dict]:
    seen = set()
    out = []
    for q in questions:
        fp = question_fingerprint(q)
        if fp in seen:
            continue
        seen.add(fp)
        out.append(q)
    return out


def is_red_hex(val: str) -> bool:
    v = (val or "").strip().lstrip("#").upper()
    if not re.fullmatch(r"[0-9A-F]{6}", v):
        return False
    r, g, b = int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16)
    return (r >= 200 and g <= 80 and b <= 80)


def is_red_run(run) -> bool:
    color = run.font.color
    if not color:
        return False
    rgb = color.rgb
    if rgb is None:
        return False
    r, g, b = rgb[0], rgb[1], rgb[2]
    return (r >= 200 and g <= 80 and b <= 80)


def paragraph_text_and_is_red(paragraph: Paragraph) -> tuple[str, bool]:
    text = "".join(run.text for run in paragraph.runs).strip()
    any_red = any(is_red_run(run) and run.text.strip() for run in paragraph.runs)
    return text, any_red


def local(tag: str) -> str:
    return (tag or "").rsplit("}", 1)[-1]


def txbx_paragraph_text_and_is_red(p_elm: CT_P) -> tuple[str, bool]:
    parts: list[str] = []
    any_red = False
    for r in p_elm.iter():
        if not str(getattr(r, "tag", "")).endswith("}r"):
            continue
        run_texts = []
        for t_node in r.iter():
            if str(getattr(t_node, "tag", "")).endswith("}t") and getattr(t_node, "text", None):
                run_texts.append(t_node.text)
        t = "".join(run_texts)
        if t:
            parts.append(t)
        if any_red or not t.strip():
            continue

        rpr = None
        for child in list(r):
            if str(getattr(child, "tag", "")).endswith("}rPr"):
                rpr = child
                break
        if rpr is None:
            continue
        color = None
        for c in rpr.iter():
            if str(getattr(c, "tag", "")).endswith("}color"):
                color = c
                break
        if color is None:
            continue
        val = None
        for k, v in getattr(color, "attrib", {}).items():
            if str(k).endswith("}val") or str(k) == "val":
                val = v
                break
        if val and is_red_hex(val):
            any_red = True

    return clean_text("".join(parts)), any_red


def textbox_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    out: list[tuple[str, bool]] = []
    txbx_paras = []
    for el in p_elm.iter():
        if str(getattr(el, "tag", "")).endswith("}txbxContent"):
            for p2 in el.iter():
                if str(getattr(p2, "tag", "")).endswith("}p"):
                    txbx_paras.append(p2)
    for tx_p in txbx_paras:
        t, red = txbx_paragraph_text_and_is_red(tx_p)
        if t:
            out.append((t, red))
    return out


def drawingml_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    lines: list[tuple[str, bool]] = []
    seen = set()
    for el in p_elm.iter():
        tag = str(getattr(el, "tag", ""))
        if not tag.endswith("}p"):
            continue
        ns = tag.split("}")[0].lstrip("{")
        if "drawing" not in ns.lower():
            continue

        parts: list[str] = []
        any_red = False
        for node in el.iter():
            ntag = str(getattr(node, "tag", ""))
            if ntag.endswith("}t") and getattr(node, "text", None):
                parts.append(node.text)
            if not any_red and ntag.endswith("}srgbClr"):
                val = None
                for k, v in getattr(node, "attrib", {}).items():
                    if str(k).endswith("}val") or str(k) == "val":
                        val = v
                        break
                if val and is_red_hex(val):
                    any_red = True
        text = clean_text("".join(parts))
        if not text:
            continue
        key = (text, any_red)
        if key in seen:
            continue
        seen.add(key)
        lines.append((text, any_red))
    return lines

def iter_block_items(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def extract_items_with_red(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    items: list[dict] = []

    def push_text(t: str, is_red: bool):
        t = clean_text(t)
        if t:
            items.append({"text": t, "is_red": is_red})

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t, red = paragraph_text_and_is_red(block)
            push_text(t, red)
            for t2, red2 in textbox_texts_in_paragraph(block):
                push_text(t2, red2)
            for t2, red2 in drawingml_texts_in_paragraph(block):
                push_text(t2, red2)
        else:
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t, red = paragraph_text_and_is_red(p)
                        push_text(t, red)
                        for t2, red2 in textbox_texts_in_paragraph(p):
                            push_text(t2, red2)
                        for t2, red2 in drawingml_texts_in_paragraph(p):
                            push_text(t2, red2)
    return items


def extract_items_with_red_v1(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    items: list[dict] = []

    def push_text(t: str, is_red: bool):
        t = clean_text(t)
        if t:
            items.append({"text": t, "is_red": bool(is_red)})

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t, red = paragraph_text_and_is_red(block)
            push_text(t, red)
        else:
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t, red = paragraph_text_and_is_red(p)
                        push_text(t, red)
    return items

NOISE_RE_V1 = re.compile(
    r"^(Instructions|For learners|For assessors|For students|Range and conditions|Decision-making rules|"
    r"Pre-approved reasonable adjustments|Rubric|Knowledge Test)\b",
    re.IGNORECASE,
)

QUESTION_CMD_INNER_RE_V1 = re.compile(
    r"\b(Which\s+of\s+the\s+following\b|" r"(Identify|Select)\s+(one|two|three|four|five|six|seven|eight|nine|ten|\d+)\b)",
    re.IGNORECASE,
)

COMMAND_QUESTION_RE_V1 = re.compile(
    r"^(Illustrate|Critically\s+(?:assess|analyse|analyze|evaluate)|"
    r"Evaluate|Determine|Articulate|Prescribe|Analyse|Analyze|Review|Recommend)\b.+",
    re.IGNORECASE,
)

RUBRIC_START_RE_V1 = re.compile(r"^Answer\s+needs\s+to\s+address\b", re.IGNORECASE)
ESSAY_GUIDE_RE_V1 = re.compile(r"^Answer\s+(may|must)\s+address", re.IGNORECASE)


def build_description_v1(items: list[dict]) -> str:
    collecting = False
    lines = []
    for it in items:
        t = clean_text(it.get("text", ""))
        if not t:
            continue
        if re.search(r"\bFor learners\b", t, re.IGNORECASE):
            collecting = True
        if collecting and (
            QUESTION_CMD_INNER_RE_V1.search(t)
            or COMMAND_QUESTION_RE_V1.match(strip_q_prefix(t))
            or re.search(r"\bdragging\s+and\s+dropping\b|\bdrag\s+and\s+drop\b|\bComplete\s+the\s+table\b", t, re.IGNORECASE)
        ):
            break
        if collecting:
            lines.append(t)
    if not lines:
        return ""

    html_parts = []
    in_list = False
    for ln in lines:
        ln = ln.strip()
        if "•" in ln and not ln.strip().startswith("•"):
            before, *bullets = [p.strip() for p in ln.split("•") if p.strip()]
            if before:
                if in_list:
                    html_parts.append("</ul>")
                    in_list = False
                html_parts.append(f"<p>{before}</p>")
            if bullets:
                if not in_list:
                    html_parts.append("<ul>")
                    in_list = True
                for b in bullets:
                    html_parts.append(f"<li>{b}</li>")
            continue
        if ln.startswith("•"):
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{ln.lstrip('•').strip()}</li>")
            continue
        if in_list:
            html_parts.append("</ul>")
            in_list = False
        html_parts.append(f"<p>{ln}</p>")
    if in_list:
        html_parts.append("</ul>")
    return "\n".join(html_parts)


def parse_mcq_questions_v1(items: list[dict]) -> list[dict]:
    questions_list = []
    current_q = None
    current_opts: list[dict] = []

    def flush():
        nonlocal current_q, current_opts
        if not current_q:
            return
        opts = [o for o in current_opts if not NOISE_RE_V1.match(o["text"])]
        option_texts = [o["text"] for o in opts]
        correct = [i for i, o in enumerate(opts) if o["is_red"]]
        qtext = strip_q_prefix(current_q.strip())
        qlower = qtext.lower()
        multi = bool(re.search(r"\bselect\s+(two|three|four|five|\d+)", qlower)) or ("apply" in qlower) or (len(correct) > 1)
        questions_list.append({"question": qtext, "options": option_texts, "correct": correct, "multi": multi, "kind": "mcq"})
        current_q = None
        current_opts = []

    for it in items:
        line = clean_text(it.get("text", ""))
        if not line:
            continue
        if NOISE_RE_V1.match(line):
            continue
        if ESSAY_GUIDE_RE_V1.match(line):
            current_q = None
            current_opts = []
            continue

        m = QUESTION_CMD_INNER_RE_V1.search(line)
        if m:
            flush()
            start = m.start()
            stem = line[:start].strip()
            cmd_plus = line[start:].strip()
            q_line = f"{stem} {cmd_plus}".strip() if stem else cmd_plus
            current_q = strip_q_prefix(q_line)
            current_opts = []
            continue

        if current_q:
            current_opts.append({"text": line, "is_red": it.get("is_red", False)})

    flush()
    return [q for q in questions_list if len(q.get("options") or []) >= 2 and len(q.get("question") or "") >= 10]


def parse_essay_questions_v1(items: list[dict]) -> list[dict]:
    questions = []
    n = len(items)
    i = 0
    while i < n:
        raw = clean_text(items[i].get("text", ""))
        if not raw or NOISE_RE_V1.match(raw):
            i += 1
            continue
        line = strip_q_prefix(raw)
        if COMMAND_QUESTION_RE_V1.match(line):
            j = i + 1
            next_line = ""
            while j < n:
                nxt = clean_text(items[j].get("text", ""))
                if nxt and not NOISE_RE_V1.match(nxt):
                    next_line = nxt
                    break
                j += 1
            if RUBRIC_START_RE_V1.match(next_line):
                questions.append({"question": line, "options": [], "correct": [], "multi": False, "kind": "essay"})
                i = j + 1
                continue
        i += 1
    return [q for q in questions if len((q.get("question") or "").strip()) >= 10]


MATCHING_STEM_RE = re.compile(
    r"\b(complete\s+the\s+table|drag(?:ging)?\s+and\s+drop(?:ping)?|drag\s+and\s+drop|match\s+each|match\s+the\s+following|match\s+.*\s+to\s+the\s+correct|select\s+one.*for\s+each)\b",
    re.IGNORECASE,
)


def looks_like_matching_stem(t: str) -> bool:
    t2 = strip_q_prefix(clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors", "for students")):
        return False
    return bool(MATCHING_STEM_RE.search(t2))


def join_lines(lines: list[str]) -> str:
    parts = [clean_text(x) for x in (lines or []) if clean_text(x)]
    return "; ".join(parts).strip()


def table_to_grid(tbl: Table) -> list[list[list[str]]]:
    grid: list[list[list[str]]] = []
    for row in tbl.rows:
        r: list[list[str]] = []
        for cell in row.cells:
            lines = []
            for p in cell.paragraphs:
                t = clean_text(p.text)
                if t:
                    lines.append(t)
            r.append(lines)
        grid.append(r)
    return grid


def table_fingerprint(grid: list[list[list[str]]]) -> str:
    rows = []
    for r in grid:
        rows.append("|".join(join_lines(c) for c in r))
    return normalize_key("||".join(rows))


def score_columns(grid, a: int, b: int) -> int:
    sc = 0
    for r in grid[1:]:
        if a >= len(r) or b >= len(r):
            continue
        left = join_lines(r[a])
        right = join_lines(r[b])
        if left and right:
            sc += 1
    return sc


def pick_best_columns(grid):
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    best = None
    best_sc = 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = score_columns(grid, a, b)
            if sc > best_sc:
                best_sc = sc
                best = (a, b)
    if best_sc < 2:
        return None
    return best


def extract_pairs(grid, left_col: int, right_col: int, start_row: int = 1):
    pairs = []
    for r in grid[start_row:]:
        if left_col >= len(r) or right_col >= len(r):
            continue
        left = join_lines(r[left_col])
        right = join_lines(r[right_col])
        if not left or not right:
            continue
        pairs.append({"left": left, "right": right})
    return pairs


def parse_matching_questions_doc_order(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    out = []
    recent: list[str] = []
    seen = set()
    seq = 0

    def choose_stem() -> str | None:
        for t in reversed(recent[-50:]):
            if looks_like_matching_stem(t):
                return strip_q_prefix(t)
        return None

    for el in iter_block_items(doc):
        seq += 1
        if isinstance(el, Paragraph):
            t, _ = paragraph_text_and_is_red(el)
            t = clean_text(t)
            if t:
                recent.append(t)
            continue

        grid = table_to_grid(el)
        fp = table_fingerprint(grid)
        if fp in seen:
            continue
        seen.add(fp)

        cols = pick_best_columns(grid)
        if not cols:
            continue
        left_col, right_col = cols
        pairs = extract_pairs(grid, left_col, right_col, start_row=1)
        if len(pairs) < 2:
            continue

        stem = choose_stem()
        if not stem:
            header = grid[0] if grid else []
            hL = (join_lines(header[left_col]) if header and left_col < len(header) else "Left")
            hR = (join_lines(header[right_col]) if header and right_col < len(header) else "Right")
            stem = f"Match each '{hL}' to the correct '{hR}'."

        out.append({"question": stem, "pairs": pairs, "kind": "matching", "options": [], "correct": [], "multi": False, "_order": seq})
        recent = []

    return out


MATCHING_STEM_RE_V1_EXACT = re.compile(
    r"\b("
    r"complete\s+the\s+table|"
    r"drag(?:ging)?\s+and\s+drop(?:ping)?|"
    r"drag\s+and\s+drop|"
    r"match\s+each|"
    r"match\s+the|"
    r"match\s+.*\s+to\s+the\s+correct|"
    r"select\s+one.*for\s+each"
    r")\b",
    re.IGNORECASE,
)

INSTRUCTION_TABLE_NOISE_RE_V1_EXACT = re.compile(
    r"\b("
    r"range\s+and\s+conditions?|decision-?making\s+rules?|"
    r"rubric|pre-?approved\s+reasonable\s+adjustments?|"
    r"for\s+learners?|for\s+assessors?|instructions?|"
    r"evidence|required|criteria|competent|nyc|submission|marking"
    r")\b",
    re.IGNORECASE,
)


def looks_like_matching_stem_v1_exact(t: str) -> bool:
    t2 = strip_q_prefix(clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors")):
        return False
    if COMMAND_QUESTION_RE_V1.match(t2):
        return False
    if "which of the following" in low:
        return False
    return bool(MATCHING_STEM_RE_V1_EXACT.search(t2))


def cell_lines_v1_exact(cell) -> list[str]:
    lines: list[str] = []
    for p in cell.paragraphs:
        t, _ = paragraph_text_and_is_red(p)
        t = clean_text(t)
        if not t:
            continue
        if "•" in t:
            parts = [x.strip() for x in t.split("•") if x.strip()]
            lines.extend(parts)
        else:
            lines.append(t)

    out: list[str] = []
    seen: set[str] = set()
    for x in lines:
        x = clean_text(x)
        if not x:
            continue
        k = normalize_key(x)
        if k in seen:
            continue
        seen.add(k)
        out.append(x)
    return out


def table_to_grid_v1_exact(table: Table) -> list[list[list[str]]]:
    return [[cell_lines_v1_exact(c) for c in row.cells] for row in table.rows]


def table_fingerprint_v1_exact(grid) -> str:
    flat: list[str] = []
    for row in grid:
        for cell in row:
            flat.append("|".join(cell))
    blob = "||".join([normalize_key(x) for x in flat if x])
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()


def is_instruction_table_v1_exact(grid) -> bool:
    texts: list[str] = []
    for row in grid:
        for cell in row:
            if cell:
                texts.append(" ".join(cell))
    blob = " ".join(texts).strip()
    if not blob:
        return True
    first_row = " ".join([join_lines(c) for c in (grid[0] if grid else [])]).lower()
    if "range and conditions" in first_row or "decision-making rules" in first_row or "pre-approved" in first_row:
        return True
    hits = sum(1 for t in texts if INSTRUCTION_TABLE_NOISE_RE_V1_EXACT.search(t))
    ratio = hits / max(1, len(texts))
    return ratio >= 0.40


def pair_is_valid_v1_exact(left: str, right: str) -> bool:
    if not left or not right:
        return False
    if normalize_key(left) == normalize_key(right):
        return False
    if len(left) > 180 or len(right) > 350:
        return False
    return True


def guess_header_skip_v1_exact(grid) -> int:
    if not grid or not grid[0]:
        return 0
    row0 = " ".join([join_lines(c) for c in grid[0] if c]).strip().lower()
    header_words = [
        "definition",
        "term",
        "meaning",
        "word",
        "concept",
        "numbers",
        "number",
        "example",
        "type",
        "classification",
        "left",
        "right",
    ]
    if any(w in row0 for w in header_words):
        return 1
    nonempty = [join_lines(c) for c in grid[0] if join_lines(c)]
    if nonempty and sum(len(x) <= 20 for x in nonempty) / len(nonempty) >= 0.8:
        return 1
    return 0


def score_columns_v1_exact(grid, a: int, b: int) -> int:
    start = guess_header_skip_v1_exact(grid)
    score = 0
    for r in range(start, len(grid)):
        if a >= len(grid[r]) or b >= len(grid[r]):
            continue
        left = join_lines(grid[r][a])
        right = join_lines(grid[r][b])
        if pair_is_valid_v1_exact(left, right):
            score += 1
    return score


def pick_best_columns_v1_exact(grid) -> tuple[int, int] | None:
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    if max_cols < 2:
        return None
    best, best_score = None, 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = score_columns_v1_exact(grid, a, b)
            if sc > best_score:
                best_score = sc
                best = (a, b)
    if best is None or best_score < 2:
        return None
    return best


def extract_pairs_v1_exact(grid, left_col: int, right_col: int, start_row: int = 0) -> list[dict]:
    pairs = []
    for r in range(start_row, len(grid)):
        if left_col >= len(grid[r]) or right_col >= len(grid[r]):
            continue
        left = join_lines(grid[r][left_col])
        right = join_lines(grid[r][right_col])
        left = re.sub(r"^\(?[a-z]\)\s*", "", left, flags=re.IGNORECASE).strip()
        left = re.sub(r"^[a-z]\.\s*", "", left, flags=re.IGNORECASE).strip()
        if not pair_is_valid_v1_exact(left, right):
            continue
        pairs.append({"left": left, "right": right})

    seen, out = set(), []
    for p in pairs:
        k = normalize_key(p["left"]) + "=>" + normalize_key(p["right"])
        if k in seen:
            continue
        seen.add(k)
        out.append(p)
    return out


def iter_elements_recursive_v1_exact(container):
    if hasattr(container, "element"):
        parent_elm = container.element.body
        parent_obj = container
    else:
        parent_elm = container._tc
        parent_obj = container
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield ("p", Paragraph(child, parent_obj))
        elif child.tag.endswith("}tbl"):
            tbl = Table(child, parent_obj)
            yield ("tbl", tbl)
            for row in tbl.rows:
                for cell in row.cells:
                    yield from iter_elements_recursive_v1_exact(cell)


def cell_fill_hex_v1_exact(cell) -> str | None:
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill:
        return None
    fill = fill.strip().upper()
    if fill in ("AUTO", "FFFFFF"):
        return None
    return fill


def row_fill_signature_v1_exact(row_cells) -> tuple:
    return tuple(cell_fill_hex_v1_exact(c) for c in row_cells)


def guess_header_skip_by_row_color_v1_exact(table: Table) -> int:
    if not table.rows:
        return 0
    row0 = table.rows[0]
    sig0 = row_fill_signature_v1_exact(row0.cells)
    if any(sig0):
        sample = table.rows[1 : min(6, len(table.rows))]
        non_header_like = 0
        for r in sample:
            sig = row_fill_signature_v1_exact(r.cells)
            if not any(sig):
                non_header_like += 1
        if non_header_like >= max(1, len(sample) - 1):
            return 1
    return 0


def column_fill_stats_v1_exact(table: Table):
    stats = {}
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            f = cell_fill_hex_v1_exact(cell)
            if f:
                stats[ci] = stats.get(ci, 0) + 1
    return stats


def pick_term_column_by_fill_v1_exact(table: Table) -> int | None:
    stats = column_fill_stats_v1_exact(table)
    if not stats:
        return None
    best_col = max(stats, key=lambda c: stats[c])
    best = stats[best_col]
    sorted_counts = sorted(stats.values(), reverse=True)
    second = sorted_counts[1] if len(sorted_counts) > 1 else 0
    if best >= 2 and best >= second + 2:
        return best_col
    return None


def parse_matching_questions_doc_order_v1_exact(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    out = []
    recent_paras: list[str] = []
    seen_tables = set()
    MAX_LOOKBACK = 50

    def choose_stem() -> str | None:
        for t in reversed(recent_paras[-MAX_LOOKBACK:]):
            if looks_like_matching_stem_v1_exact(t):
                return strip_q_prefix(clean_text(t))
        return None

    for kind, el in iter_elements_recursive_v1_exact(doc):
        if kind == "p":
            t, _ = paragraph_text_and_is_red(el)
            t = clean_text(t)
            if t:
                recent_paras.append(t)
                if len(recent_paras) > 400:
                    recent_paras = recent_paras[-400:]
            continue

        grid = table_to_grid_v1_exact(el)
        header_skip = guess_header_skip_by_row_color_v1_exact(el)

        tfp = table_fingerprint_v1_exact(grid)
        if tfp in seen_tables:
            continue
        seen_tables.add(tfp)

        if is_instruction_table_v1_exact(grid):
            continue

        term_col = pick_term_column_by_fill_v1_exact(el)
        if term_col is not None:
            max_cols = max(len(r) for r in grid)
            best_right = None
            best_score = 0
            for b in range(max_cols):
                if b == term_col:
                    continue
                sc = score_columns_v1_exact(grid, term_col, b)
                if sc > best_score:
                    best_score = sc
                    best_right = b
            if best_right is None or best_score < 2:
                continue
            left_col, right_col = term_col, best_right
        else:
            cols = pick_best_columns_v1_exact(grid)
            if not cols:
                continue
            left_col, right_col = cols

        pairs = extract_pairs_v1_exact(grid, left_col, right_col, start_row=header_skip)
        if len(pairs) < 2:
            continue

        stem = choose_stem()
        if not stem:
            header = grid[0] if grid else []
            hL = (header[left_col][0] if header and left_col < len(header) and header[left_col] else "Left")
            hR = (header[right_col][0] if header and right_col < len(header) and header[right_col] else "Right")
            stem = f"Match each '{hL}' to the correct '{hR}'."

        out.append({"question": stem, "pairs": pairs, "kind": "matching", "options": [], "correct": [], "multi": False})
        recent_paras = []

    return out

V2_Q_PREFIX_RE = re.compile(r"^\s*(?:lo\s*)?(?:question|q)\s*\d+\s*[\.\)]\s*", re.IGNORECASE)
V2_ANSWER_GUIDE_INLINE_RE = re.compile(r"\bAnswer\s+(?:may|must|needs)\s+address\b[:\s-]*", re.IGNORECASE)
V2_ANSWER_GUIDE_START_RE = re.compile(r"^\s*Answer\s+(?:may|must|needs)\s+address\b", re.IGNORECASE)
V2_HARD_QNUM_RE = re.compile(r"(?:(?<=\s)|^)(?:lo\s*)?(?:question|q)\s*\d+\s*[\.\)]\s*", re.IGNORECASE)
V2_HARD_QNUM_RANGE_RE = re.compile(r"(?:(?<=\s)|^)(?:lo\s*)?q\s*\d+\s*[-–]\s*\d+\s*[\.\)]\s*", re.IGNORECASE)
V2_MATCHING_STEM_RE = re.compile(
    r"\b(complete\s+the\s+table|drag(?:ging)?\s+and\s+drop(?:ping)?|drag\s+and\s+drop|match\s+each|match\s+the\s+following|match\s+.*\s+to\s+the\s+correct|select\s+one.*for\s+each)\b",
    re.IGNORECASE,
)


def v2_clean_text(s: str) -> str:
    s = (s or "").replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def v2_normalize_key(s: str) -> str:
    return v2_clean_text(s).lower()


def v2_strip_q_prefix(s: str) -> str:
    return v2_clean_text(V2_Q_PREFIX_RE.sub("", s or "", count=1))


def v2_strip_answer_guide(text: str) -> str:
    t = v2_clean_text(text)
    if not t:
        return ""
    m = V2_ANSWER_GUIDE_INLINE_RE.search(t)
    if not m:
        return t
    return v2_clean_text(t[: m.start()])


def v2_trim_after_question_mark(text: str) -> str:
    t = v2_clean_text(text)
    if not t:
        return ""
    if "?" not in t:
        return t
    qpos = t.find("?")
    return v2_clean_text(t[: qpos + 1])


def v2_txbx_paragraph_text_and_is_red(p_elm: CT_P) -> tuple[str, bool]:
    parts: list[str] = []
    any_red = False
    for r in p_elm.iter():
        if not str(getattr(r, "tag", "")).endswith("}r"):
            continue
        run_texts = []
        for t_node in r.iter():
            if str(getattr(t_node, "tag", "")).endswith("}t") and getattr(t_node, "text", None):
                run_texts.append(t_node.text)
        t = "".join(run_texts)
        if t:
            parts.append(t)
        if any_red or not t.strip():
            continue

        rpr = None
        for child in list(r):
            if str(getattr(child, "tag", "")).endswith("}rPr"):
                rpr = child
                break
        if rpr is None:
            continue

        color = None
        for c in rpr.iter():
            if str(getattr(c, "tag", "")).endswith("}color"):
                color = c
                break
        if color is None:
            continue
        val = None
        for k, v in getattr(color, "attrib", {}).items():
            if str(k).endswith("}val") or str(k) == "val":
                val = v
                break
        if val and is_red_hex(val):
            any_red = True

    return v2_clean_text("".join(parts)), any_red


def v2_textbox_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    out: list[tuple[str, bool]] = []
    txbx_paras = []
    for el in p_elm.iter():
        if str(getattr(el, "tag", "")).endswith("}txbxContent"):
            for p2 in el.iter():
                if str(getattr(p2, "tag", "")).endswith("}p"):
                    txbx_paras.append(p2)
    for tx_p in txbx_paras:
        t, red = v2_txbx_paragraph_text_and_is_red(tx_p)
        if t:
            out.append((t, red))
    return out


def v2_drawingml_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    lines: list[tuple[str, bool]] = []
    seen = set()
    for el in p_elm.iter():
        tag = str(getattr(el, "tag", ""))
        if not tag.endswith("}p"):
            continue
        ns = tag.split("}")[0].lstrip("{")
        if "drawing" not in ns.lower():
            continue

        parts: list[str] = []
        any_red = False
        for node in el.iter():
            ntag = str(getattr(node, "tag", ""))
            if ntag.endswith("}t") and getattr(node, "text", None):
                parts.append(node.text)
            if not any_red and ntag.endswith("}srgbClr"):
                val = None
                for k, v in getattr(node, "attrib", {}).items():
                    if str(k).endswith("}val") or str(k) == "val":
                        val = v
                        break
                if val and is_red_hex(val):
                    any_red = True
        text = v2_clean_text("".join(parts))
        if not text:
            continue
        key = (text, any_red)
        if key in seen:
            continue
        seen.add(key)
        lines.append((text, any_red))
    return lines


def v2_extract_items_with_red(docx_path: str) -> list[dict]:
    """
    Exact extractor from canvas_ai_app.py.
    """

    def extract_with_python_docx(path: str) -> list[dict]:
        doc = Document(path)
        items: list[dict] = []
        last = None

        def push(t: str, red: bool):
            nonlocal last
            t = v2_clean_text(t)
            if not t:
                return
            key = (t, bool(red))
            if last == key:
                return
            items.append({"text": t, "is_red": bool(red)})
            last = key

        def push_paragraph(p: Paragraph):
            t, red = paragraph_text_and_is_red(p)
            push(t, red)
            for t2, red2 in v2_textbox_texts_in_paragraph(p):
                push(t2, red2)
            for t2, red2 in v2_drawingml_texts_in_paragraph(p):
                push(t2, red2)

        def push_table(tbl: Table):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        push_paragraph(p)
                    for sub in getattr(cell, "tables", []) or []:
                        push_table(sub)

        doc_obj = doc
        for block in iter_block_items(doc_obj):
            if isinstance(block, Paragraph):
                push_paragraph(block)
            else:
                push_table(block)
        return items

    def extract_with_xml_fallback(path: str) -> list[dict]:
        try:
            with zipfile.ZipFile(path) as z:
                raw = z.read("word/document.xml")
        except Exception:
            return []
        try:
            root = ET.fromstring(raw)
        except Exception:
            return []

        def run_is_red(run_el) -> bool:
            for el in run_el.iter():
                lname = local(getattr(el, "tag", ""))
                if lname == "color":
                    for k, v in getattr(el, "attrib", {}).items():
                        if str(k).endswith("}val") or str(k) == "val":
                            if v and is_red_hex(v):
                                return True
                if lname == "srgbClr":
                    for k, v in getattr(el, "attrib", {}).items():
                        if str(k).endswith("}val") or str(k) == "val":
                            if v and is_red_hex(v):
                                return True
            return False

        items: list[dict] = []
        last = None

        def push(parts: list[str], red: bool):
            nonlocal last
            t = v2_clean_text("".join(parts))
            if not t:
                return
            key = (t, bool(red))
            if last == key:
                return
            items.append({"text": t, "is_red": bool(red)})
            last = key

        def walk(node, in_p: bool = False, buf: list[str] | None = None, red_any: list[bool] | None = None):
            lname = local(str(getattr(node, "tag", "")))
            if lname == "p":
                b: list[str] = []
                r = [False]
                for ch in list(node):
                    walk(ch, in_p=True, buf=b, red_any=r)
                push(b, r[0])
                return
            if lname == "r":
                is_red = run_is_red(node)
                for ch in list(node):
                    walk(ch, in_p=in_p, buf=buf, red_any=red_any)
                if is_red and red_any is not None:
                    red_any[0] = True
                return
            if lname == "t":
                if in_p and buf is not None and getattr(node, "text", None):
                    buf.append(node.text)
                return
            for ch in list(node):
                walk(ch, in_p=in_p, buf=buf, red_any=red_any)

        for ch in list(root):
            walk(ch)
        return items

    def score(items: list[dict]) -> int:
        if not items:
            return 0
        qverb = re.compile(
            r"^(q\s*\d+\s*[\.)]\s*)?(list|describe|explain|outline|state|name|provide|define|identify|select|choose|pick|match|complete)\b",
            re.IGNORECASE,
        )
        sc = 0
        for it in items:
            t = v2_clean_text(it.get("text", ""))
            if not t:
                continue
            if qverb.match(t) or "Which of the following" in t:
                sc += 3
            if t.endswith("?"):
                sc += 1
        return sc + min(80, len(items) // 8)

    items_docx = extract_with_python_docx(docx_path)
    items_xml = extract_with_xml_fallback(docx_path)
    return items_xml if score(items_xml) > score(items_docx) else items_docx


def v2_split_items_on_internal_qnums(items: list[dict]) -> list[dict]:
    out: list[dict] = []
    for it in items:
        t = v2_clean_text(it.get("text", ""))
        if not t:
            continue
        if V2_HARD_QNUM_RANGE_RE.search(t):
            out.append({"text": t, "is_red": bool(it.get("is_red"))})
            continue
        starts = [m.start() for m in V2_HARD_QNUM_RE.finditer(t)]
        if len(starts) <= 1:
            out.append({"text": t, "is_red": bool(it.get("is_red"))})
            continue
        for a, b in zip(starts, starts[1:] + [len(t)]):
            seg = v2_clean_text(t[a:b])
            if seg:
                out.append({"text": seg, "is_red": bool(it.get("is_red"))})
    return out


def v2_looks_like_matching_stem(t: str) -> bool:
    t2 = v2_strip_q_prefix(v2_clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors", "for students")):
        return False
    return bool(V2_MATCHING_STEM_RE.search(t2))


def v2_join_lines(lines: list[str]) -> str:
    parts = [v2_clean_text(x) for x in (lines or []) if v2_clean_text(x)]
    return "; ".join(parts).strip()


def v2_table_to_grid(tbl: Table) -> list[list[list[str]]]:
    grid: list[list[list[str]]] = []
    for row in tbl.rows:
        r: list[list[str]] = []
        for cell in row.cells:
            lines = []
            for p in cell.paragraphs:
                t = v2_clean_text(p.text)
                if t:
                    lines.append(t)
            r.append(lines)
        grid.append(r)
    return grid


def v2_table_fingerprint(grid: list[list[list[str]]]) -> str:
    rows = []
    for r in grid:
        rows.append("|".join(v2_join_lines(c) for c in r))
    return v2_normalize_key("||".join(rows))


def v2_score_columns(grid, a: int, b: int) -> int:
    sc = 0
    for r in grid[1:]:
        if a >= len(r) or b >= len(r):
            continue
        left = v2_join_lines(r[a])
        right = v2_join_lines(r[b])
        if left and right:
            sc += 1
    return sc


def v2_pick_best_columns(grid):
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    best = None
    best_sc = 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = v2_score_columns(grid, a, b)
            if sc > best_sc:
                best_sc = sc
                best = (a, b)
    if best_sc < 2:
        return None
    return best


def v2_extract_pairs(grid, left_col: int, right_col: int, start_row: int = 1):
    pairs = []
    for r in grid[start_row:]:
        if left_col >= len(r) or right_col >= len(r):
            continue
        left = v2_join_lines(r[left_col])
        right = v2_join_lines(r[right_col])
        if not left or not right:
            continue
        pairs.append({"left": left, "right": right})
    return pairs


def v2_parse_matching_questions_doc_order(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    out = []
    recent: list[str] = []
    seen = set()
    seq = 0

    def choose_stem() -> str | None:
        for t in reversed(recent[-50:]):
            if v2_looks_like_matching_stem(t):
                return v2_strip_q_prefix(t)
        return None

    for el in iter_block_items(doc):
        seq += 1
        if isinstance(el, Paragraph):
            t, _ = paragraph_text_and_is_red(el)
            t = v2_clean_text(t)
            if t:
                recent.append(t)
            continue

        grid = v2_table_to_grid(el)
        fp = v2_table_fingerprint(grid)
        if fp in seen:
            continue
        seen.add(fp)

        cols = v2_pick_best_columns(grid)
        if not cols:
            continue
        left_col, right_col = cols
        pairs = v2_extract_pairs(grid, left_col, right_col, start_row=1)
        if len(pairs) < 2:
            continue

        stem = choose_stem()
        if not stem:
            header = grid[0] if grid else []
            hL = (v2_join_lines(header[left_col]) if header and left_col < len(header) else "Left")
            hR = (v2_join_lines(header[right_col]) if header and right_col < len(header) else "Right")
            stem = f"Match each '{hL}' to the correct '{hR}'."

        out.append(
            {
                "question": stem,
                "pairs": pairs,
                "kind": "matching",
                "options": [],
                "correct": [],
                "multi": False,
                "_order": seq,
                "qnum": None,
            }
        )
        recent = []

    return out


def v2_ai_segment_items_openai(items: list[dict], cfg: OpenAIConfig) -> tuple[list[dict], list[str]]:
    log: list[str] = []
    if not items:
        return [], log

    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "questions": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "kind": {"type": "string", "enum": ["mcq", "essay"]},
                        "stem": {"type": "array", "items": {"type": "integer"}},
                        "options": {"type": "array", "items": {"type": "array", "items": {"type": "integer"}}},
                    },
                    "required": ["kind", "stem", "options"],
                },
            }
        },
        "required": ["questions"],
    }

    def to_line(i: int) -> str:
        t = v2_clean_text(items[i].get("text", ""))
        red = "R1" if items[i].get("is_red") else "R0"
        return f"I{i}|{red}|{t}"

    max_block_items = 170
    overlap = 50
    blocks: list[tuple[int, int]] = []
    start = 0
    while start < len(items):
        end = min(len(items), start + max_block_items)
        blocks.append((start, end))
        if end >= len(items):
            break
        start = max(0, end - overlap)

    all_qs: list[dict] = []
    for (a, b) in blocks:
        ctx = []
        for i in range(a, b):
            t = v2_clean_text(items[i].get("text", ""))
            if not t:
                continue
            ctx.append(to_line(i))
        if len(ctx) < 6:
            continue

        log.append(f"AI block: {a}-{b} lines={len(ctx)}")
        prompt = (
            "You are segmenting a DOCX extraction into Canvas quiz questions.\n"
            "Return STRICT JSON only (per schema).\n"
            "\n"
            "Hard rules:\n"
            "- You MUST NOT invent any text.\n"
            "- You may ONLY reference item indices (I<n>) from the provided list.\n"
            "- Keep original order (earlier indices first).\n"
            "- Do NOT create questions from instructions/policy/rubric.\n"
            "- Do NOT include assessor guide content like 'Answer may/must/needs address' or sample answers in stems/options.\n"
            "- Correct MCQ options are those with R1 (red). For essay questions: options must be [].\n"
            "\n"
            "MCQ rules:\n"
            "- Options are typically lettered (a), (b), etc or separate lines under a prompt.\n"
            "- If you cannot find at least 2 options, do NOT output an MCQ.\n"
            "\n"
            "Essay rules:\n"
            "- Use the question prompt only. Do not append long paragraphs after '?'.\n"
            "\n"
            "Items (format: I<index>|R0/R1|text):\n"
            + "\n".join(ctx)
        )

        data, err = openai_responses_json_schema(prompt, "segment_questions", schema, cfg)
        if err:
            log.append(f"  block failed: {err}")
            continue
        qs = data.get("questions") if isinstance(data, dict) else None
        if not isinstance(qs, list):
            log.append("  block skipped: missing questions[]")
            continue

        for q in qs:
            if not isinstance(q, dict):
                continue
            kind = (q.get("kind") or "").strip().lower()
            stem_ids = q.get("stem") if isinstance(q.get("stem"), list) else []
            if kind not in ("mcq", "essay"):
                continue
            if not stem_ids or not all(isinstance(x, int) for x in stem_ids):
                continue
            if any(x < 0 or x >= len(items) for x in stem_ids):
                continue

            stem_text = v2_clean_text(" ".join(v2_clean_text(items[x].get("text", "")) for x in stem_ids))
            stem_text = v2_strip_q_prefix(v2_strip_answer_guide(stem_text))
            stem_text = v2_trim_after_question_mark(stem_text)
            if not stem_text or len(stem_text) < 10:
                continue
            if stem_text.lower().startswith(("answer may address", "answer must address", "answer needs to address")):
                continue

            if kind == "essay":
                all_qs.append(
                    {
                        "question": stem_text,
                        "options": [],
                        "correct": [],
                        "multi": False,
                        "kind": "essay",
                        "_order": min(stem_ids),
                        "qnum": None,
                    }
                )
                continue

            opt_groups = q.get("options") if isinstance(q.get("options"), list) else []
            option_texts: list[str] = []
            correct: list[int] = []
            for group in opt_groups:
                if not isinstance(group, list) or not group or not all(isinstance(x, int) for x in group):
                    continue
                if any(x < 0 or x >= len(items) for x in group):
                    continue
                t = v2_clean_text(" ".join(v2_clean_text(items[x].get("text", "")) for x in group))
                if not t:
                    continue
                if V2_ANSWER_GUIDE_START_RE.match(t):
                    continue
                option_texts.append(t)
                if any(bool(items[x].get("is_red")) for x in group):
                    correct.append(len(option_texts) - 1)

            # de-dupe options keep order
            seen = set()
            out_opts = []
            out_corr = []
            for i_opt, opt in enumerate(option_texts):
                k = v2_normalize_key(opt)
                if k in seen:
                    continue
                seen.add(k)
                if i_opt in correct:
                    out_corr.append(len(out_opts))
                out_opts.append(opt)
            if len(out_opts) < 2:
                continue

            all_qs.append(
                {
                    "question": stem_text,
                    "options": out_opts,
                    "correct": out_corr,
                    "multi": ("apply" in stem_text.lower()) or (len(out_corr) > 1),
                    "kind": "mcq",
                    "_order": min(stem_ids),
                    "qnum": None,
                }
            )

    all_qs.sort(key=lambda q: int(q.get("_order", 10**9)))
    return all_qs, log


# ===================================================
# v3 engine (EXACT) — copied from /Users/jargalmaa/Downloads/web_ui/backups/ai 43.py (BUILD_ID ai-app-2026-02-08.4)
# ===================================================
V3_Q_PREFIX_RE = re.compile(r"^\s*(?:lo\s*)?(?:question|q)\s*\d+\s*[\.\)]\s*", re.IGNORECASE)
V3_ANSWER_GUIDE_INLINE_RE = re.compile(
    r"\bAnswer\s+(?:may|must|need(?:s)?)\s+(?:to\s+)?address\b[:\s-]*",
    re.IGNORECASE,
)
V3_ANSWER_GUIDE_START_RE = re.compile(
    r"^\s*Answer\s+(?:may|must|need(?:s)?)\s+(?:to\s+)?address\b",
    re.IGNORECASE,
)
V3_ANSWER_GUIDE_ANY_RE = re.compile(
    r"\bAnswer\s+(?:may|must|need(?:s)?)\s+(?:to\s+)?address\b",
    re.IGNORECASE,
)
V3_HARD_QNUM_RE = re.compile(r"(?:(?<=\s)|^)(?:lo\s*)?(?:question|q)\s*\d+\s*[\.\)]\s*", re.IGNORECASE)
V3_HARD_QNUM_RANGE_RE = re.compile(r"(?:(?<=\s)|^)(?:lo\s*)?q\s*\d+\s*[-–]\s*\d+\s*[\.\)]\s*", re.IGNORECASE)
V3_MATCHING_STEM_RE = re.compile(
    r"\b(complete\s+the\s+table|drag(?:ging)?\s+and\s+drop(?:ping)?|drag\s+and\s+drop|match\s+each|match\s+the\s+following|match\s+.*\s+to\s+the\s+correct|select\s+one.*for\s+each)\b",
    re.IGNORECASE,
)


def v3_clean_text(s: str) -> str:
    s = (s or "").replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def v3_normalize_key(s: str) -> str:
    return v3_clean_text(s).lower()


def v3_strip_q_prefix(s: str) -> str:
    return v3_clean_text(V3_Q_PREFIX_RE.sub("", s or "", count=1))


def v3_strip_answer_guide(text: str) -> str:
    t = v3_clean_text(text)
    if not t:
        return ""
    m = V3_ANSWER_GUIDE_INLINE_RE.search(t)
    if not m:
        return t
    return v3_clean_text(t[: m.start()])


def v3_trim_after_question_mark(text: str) -> str:
    t = v3_clean_text(text)
    if not t:
        return ""
    if "?" not in t:
        return t
    qpos = t.find("?")
    return v3_clean_text(t[: qpos + 1])


def v3_trim_after_sentence_if_long(text: str, max_chars: int = 220) -> str:
    t = v3_clean_text(text)
    if len(t) <= max_chars:
        return t
    for sep in [". ", "; ", " - "]:
        pos = t.find(sep)
        if 20 <= pos <= max_chars:
            return v3_clean_text(t[: pos + (1 if sep.startswith(".") else 0)])
    return v3_clean_text(t[:max_chars])


def v3_txbx_paragraph_text_and_is_red(p_elm: CT_P) -> tuple[str, bool]:
    parts: list[str] = []
    any_red = False
    for r in p_elm.iter():
        if not str(getattr(r, "tag", "")).endswith("}r"):
            continue
        run_texts = []
        for t_node in r.iter():
            if str(getattr(t_node, "tag", "")).endswith("}t") and getattr(t_node, "text", None):
                run_texts.append(t_node.text)
        t = "".join(run_texts)
        if t:
            parts.append(t)
        if any_red or not t.strip():
            continue

        rpr = None
        for child in list(r):
            if str(getattr(child, "tag", "")).endswith("}rPr"):
                rpr = child
                break
        if rpr is None:
            continue

        color = None
        for c in rpr.iter():
            if str(getattr(c, "tag", "")).endswith("}color"):
                color = c
                break
        if color is None:
            continue
        val = None
        for k, v in getattr(color, "attrib", {}).items():
            if str(k).endswith("}val") or str(k) == "val":
                val = v
                break
        if val and is_red_hex(val):
            any_red = True

    return v3_clean_text("".join(parts)), any_red


def v3_textbox_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    out: list[tuple[str, bool]] = []
    txbx_paras = []
    for el in p_elm.iter():
        if str(getattr(el, "tag", "")).endswith("}txbxContent"):
            for p2 in el.iter():
                if str(getattr(p2, "tag", "")).endswith("}p"):
                    txbx_paras.append(p2)
    for tx_p in txbx_paras:
        t, red = v3_txbx_paragraph_text_and_is_red(tx_p)
        if t:
            out.append((t, red))
    return out


def v3_drawingml_texts_in_paragraph(paragraph: Paragraph) -> list[tuple[str, bool]]:
    try:
        p_elm = paragraph._p
    except Exception:
        return []

    lines: list[tuple[str, bool]] = []
    seen = set()
    for el in p_elm.iter():
        tag = str(getattr(el, "tag", ""))
        if not tag.endswith("}p"):
            continue
        ns = tag.split("}")[0].lstrip("{")
        if "drawing" not in ns.lower():
            continue

        parts: list[str] = []
        any_red = False
        for node in el.iter():
            ntag = str(getattr(node, "tag", ""))
            if ntag.endswith("}t") and getattr(node, "text", None):
                parts.append(node.text)
            if not any_red and ntag.endswith("}srgbClr"):
                val = None
                for k, v in getattr(node, "attrib", {}).items():
                    if str(k).endswith("}val") or str(k) == "val":
                        val = v
                        break
                if val and is_red_hex(val):
                    any_red = True
        text = v3_clean_text("".join(parts))
        if not text:
            continue
        key = (text, any_red)
        if key in seen:
            continue
        seen.add(key)
        lines.append((text, any_red))
    return lines


def v3_extract_items_with_red(docx_path: str, include_tables: bool = True) -> list[dict]:
    """
    Exact extractor from ai 43.py.
    - Prefers python-docx extraction (keeps src='body'|'table')
    - No XML fallback
    """

    def extract_with_python_docx(path: str) -> list[dict]:
        doc = Document(path)
        items: list[dict] = []
        last = None

        def push(t: str, red: bool, src: str):
            nonlocal last
            t = v3_clean_text(t)
            if not t:
                return
            key = (t, bool(red), src)
            if last == key:
                return
            items.append({"text": t, "is_red": bool(red), "src": src})
            last = key

        def push_paragraph(p: Paragraph, src: str):
            t, red = paragraph_text_and_is_red(p)
            push(t, red, src)
            for t2, red2 in v3_textbox_texts_in_paragraph(p):
                push(t2, red2, src)
            for t2, red2 in v3_drawingml_texts_in_paragraph(p):
                push(t2, red2, src)

        def push_table(tbl: Table):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        push_paragraph(p, "table")
                    for sub in getattr(cell, "tables", []) or []:
                        push_table(sub)

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                push_paragraph(block, "body")
            else:
                if include_tables:
                    push_table(block)
        return items

    return extract_with_python_docx(docx_path)


def v3_split_items_on_internal_qnums(items: list[dict]) -> list[dict]:
    out: list[dict] = []
    for it in items:
        t = v3_clean_text(it.get("text", ""))
        if not t:
            continue
        src = it.get("src")
        if V3_HARD_QNUM_RANGE_RE.search(t):
            out.append({"text": t, "is_red": bool(it.get("is_red")), "src": src})
            continue
        starts = [m.start() for m in V3_HARD_QNUM_RE.finditer(t)]
        if len(starts) <= 1:
            out.append({"text": t, "is_red": bool(it.get("is_red")), "src": src})
            continue
        for a, b in zip(starts, starts[1:] + [len(t)]):
            seg = v3_clean_text(t[a:b])
            if seg:
                out.append({"text": seg, "is_red": bool(it.get("is_red")), "src": src})
    return out


def v3_join_lines(lines: list[str]) -> str:
    parts = [v3_clean_text(x) for x in (lines or []) if v3_clean_text(x)]
    return "; ".join(parts).strip()


def v3_table_to_grid(tbl: Table) -> list[list[list[str]]]:
    grid: list[list[list[str]]] = []
    for row in tbl.rows:
        r: list[list[str]] = []
        for cell in row.cells:
            lines = []
            for p in cell.paragraphs:
                t = v3_clean_text(p.text)
                if t:
                    lines.append(t)
            r.append(lines)
        grid.append(r)
    return grid


def v3_header_contains(grid: list[list[list[str]]], *needles: str) -> bool:
    if not grid or not grid[0]:
        return False
    header = " | ".join(v3_join_lines(c) for c in grid[0]).lower()
    return all(n.lower() in header for n in needles)


def v3_iter_tables_recursive(tbl: Table):
    yield tbl
    for row in tbl.rows:
        for cell in row.cells:
            for sub in getattr(cell, "tables", []) or []:
                yield from v3_iter_tables_recursive(sub)


def v3_iter_all_tables(doc: Document):
    for el in iter_block_items(doc):
        if isinstance(el, Table):
            yield from v3_iter_tables_recursive(el)


def v3_find_item_index(items: list[dict], needle: str) -> int | None:
    n = v3_normalize_key(needle)
    if not n:
        return None
    for i, it in enumerate(items):
        t = v3_normalize_key(it.get("text", ""))
        if not t:
            continue
        if t == n or n in t:
            return i
    return None


def v3_parse_table_defined_terms_as_essays(docx_path: str, items: list[dict]) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    seen_terms: set[str] = set()
    for el in v3_iter_all_tables(doc):
        grid = v3_table_to_grid(el)
        if not v3_header_contains(grid, "poultry ingredient", "definition", "style"):
            continue

        for r in grid[1:]:
            if not r:
                continue
            term = v3_join_lines(r[0] if len(r) > 0 else [])
            term = v3_strip_q_prefix(term)
            term = v3_clean_text(re.sub(r"^\s*(?:q\s*)?\d+\s*[\.\)]\s*", "", term, flags=re.IGNORECASE))
            if not term:
                continue
            if term.lower() in {"poultry ingredient", "definition", "style/method of cooking"}:
                continue
            k = v3_normalize_key(term)
            if not k or k in seen_terms:
                continue
            seen_terms.add(k)

            order = v3_find_item_index(items, term) or v3_find_item_index(items, f"Define {term}") or 10**9
            q = f"Define: {term}. Provide one style/method of cooking."
            out.append({"question": q, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": order, "qnum": None})
    return out


def v3_parse_table_characteristics_as_essays(docx_path: str, items: list[dict]) -> list[dict]:
    doc = Document(docx_path)
    out: list[dict] = []
    seen_terms: set[str] = set()
    for el in v3_iter_all_tables(doc):
        grid = v3_table_to_grid(el)
        if not v3_header_contains(grid, "poultry type", "essential"):
            continue

        for r in grid[1:]:
            if not r:
                continue
            term = v3_join_lines(r[0] if len(r) > 0 else [])
            term = v3_strip_q_prefix(term)
            term = v3_clean_text(re.sub(r"^\s*(?:q\s*)?\d+\s*[\.\)]\s*", "", term, flags=re.IGNORECASE))
            if not term:
                continue
            if term.lower() in {"poultry type or cut", "essential characteristics"}:
                continue
            k = v3_normalize_key(term)
            if not k or k in seen_terms:
                continue
            seen_terms.add(k)
            q = f"Describe the essential characteristics of: {term}."
            order = v3_find_item_index(items, term) or v3_find_item_index(items, f"Describe {term}") or 10**9
            out.append({"question": q, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": order, "qnum": None})
    return out


def v3_collect_ignore_texts_from_forced_tables(docx_path: str) -> set[str]:
    doc = Document(docx_path)
    ignore: set[str] = set()

    def add_lines(lines: list[str]):
        for ln in lines or []:
            t = v3_clean_text(ln)
            if not t:
                continue
            if len(t) < 12 and " " not in t:
                continue
            if re.fullmatch(r"\d+", t):
                continue
            ignore.add(t)

    for tbl in v3_iter_all_tables(doc):
        grid = v3_table_to_grid(tbl)
        if v3_header_contains(grid, "poultry ingredient", "definition", "style"):
            for row in grid[1:]:
                for cell in row[1:]:
                    add_lines(cell)
            continue
        if v3_header_contains(grid, "poultry type", "essential"):
            for row in grid[1:]:
                for cell in row[1:]:
                    add_lines(cell)
            continue
        if v3_header_contains(grid, "classical chicken dishes", "contemporary chicken dishes"):
            for row in grid[1:]:
                for cell in row:
                    add_lines(cell)
            continue

    return ignore


def v3_looks_like_matching_stem(t: str) -> bool:
    t2 = v3_strip_q_prefix(v3_clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors", "for students")):
        return False
    return bool(V3_MATCHING_STEM_RE.search(t2))


def v3_table_fingerprint(grid: list[list[list[str]]]) -> str:
    rows = []
    for r in grid:
        rows.append("|".join(v3_join_lines(c) for c in r))
    return v3_normalize_key("||".join(rows))


def v3_score_columns(grid, a: int, b: int) -> int:
    sc = 0
    for r in grid[1:]:
        if a >= len(r) or b >= len(r):
            continue
        left = v3_join_lines(r[a])
        right = v3_join_lines(r[b])
        if left and right:
            sc += 1
    return sc


def v3_pick_best_columns(grid):
    if not grid:
        return None
    max_cols = max(len(r) for r in grid)
    best = None
    best_sc = 0
    for a in range(max_cols):
        for b in range(max_cols):
            if a == b:
                continue
            sc = v3_score_columns(grid, a, b)
            if sc > best_sc:
                best_sc = sc
                best = (a, b)
    if best_sc < 2:
        return None
    return best


def v3_extract_pairs(grid, left_col: int, right_col: int, start_row: int = 1):
    pairs = []
    for r in grid[start_row:]:
        if left_col >= len(r) or right_col >= len(r):
            continue
        left = v3_join_lines(r[left_col])
        right = v3_join_lines(r[right_col])
        if not left or not right:
            continue
        pairs.append({"left": left, "right": right})
    return pairs


def v3_is_table_forced_essay(grid: list[list[list[str]]]) -> bool:
    return (
        v3_header_contains(grid, "poultry ingredient", "definition", "style")
        or v3_header_contains(grid, "poultry type", "essential")
        or v3_header_contains(grid, "classical chicken dishes", "contemporary chicken dishes")
    )


def v3_parse_matching_questions_doc_order(docx_path: str, items: list[dict] | None = None) -> list[dict]:
    doc = Document(docx_path)
    out = []
    seen = set()
    seq = 0

    def is_instructions_matching(pairs: list[dict], stem: str) -> bool:
        s = v3_normalize_key(stem or "")
        if "instructions" in s and ("for students" in s or "for learners" in s or "for assessors" in s):
            return True

        left_keys = {v3_normalize_key((p.get("left") or "")) for p in (pairs or [])}
        common = {
            "range and conditions",
            "decision-making rules",
            "decision making rules",
            "pre-approved reasonable adjustments",
            "pre approved reasonable adjustments",
            "rubric",
            "instructions",
        }
        hit = sum(1 for k in common if k in left_keys)
        if hit >= 3:
            return True

        right_blob = v3_normalize_key("; ".join(p.get("right") or "" for p in (pairs or [])))
        if "students must work through this assessment independently" in right_blob:
            return True
        if "false declarations may lead to withdrawal" in right_blob:
            return True
        if "feedback comments must be provided" in right_blob:
            return True
        return False

    for el in v3_iter_all_tables(doc):
        seq += 1
        grid = v3_table_to_grid(el)
        if v3_is_table_forced_essay(grid):
            continue

        fp = v3_table_fingerprint(grid)
        if fp in seen:
            continue
        seen.add(fp)

        cols = v3_pick_best_columns(grid)
        if not cols:
            continue
        left_col, right_col = cols
        pairs = v3_extract_pairs(grid, left_col, right_col, start_row=1)
        if len(pairs) < 2:
            continue

        header = grid[0] if grid else []
        hL = (v3_join_lines(header[left_col]) if header and left_col < len(header) else "Left")
        hR = (v3_join_lines(header[right_col]) if header and right_col < len(header) else "Right")
        stem = f"Match each '{hL}' to the correct '{hR}'."

        if is_instructions_matching(pairs, stem):
            continue

        order = seq
        if items:
            for cand in (hL, hR, pairs[0].get("left", ""), pairs[0].get("right", "")):
                idx = v3_find_item_index(items, cand)
                if idx is not None:
                    order = idx
                    break

        out.append({"question": stem, "pairs": pairs, "kind": "matching", "options": [], "correct": [], "multi": False, "_order": order, "qnum": None})

    return out


V3_IGNORE_LINE_RE = re.compile(
    r"^(instructions|for learners|for assessors|for students|range and conditions|decision-making rules|pre-approved|rubric|feedback|knowledge test)\b",
    re.IGNORECASE,
)
V3_IGNORE_SECTION_RE = re.compile(
    r"^(?:range and conditions|decision-making rules|pre-approved reasonable adjustments|rubric)\b",
    re.IGNORECASE,
)
V3_IGNORE_TABLE_RE = re.compile(
    r"^(?:poultry ingredient|definition|style/method of cooking|poultry type or cut|essential characteristics|classical chicken dishes|contemporary chicken dishes)\b",
    re.IGNORECASE,
)
V3_COOKERY_METHOD_WORD_RE = re.compile(
    r"^(?:pan[-\s]?fry|deep[-\s]?fry|stir[-\s]?fry|roast|bake|grill|bbq|braise|stew|simmer|poach|saute|sauté|steam|boil)\b",
    re.IGNORECASE,
)
V3_QUESTION_START_RE = re.compile(
    r"^(?:"
    r"(?:lo\s*)?(?:question|q)\s*\d+\s*[\.\)]\s*"
    r")?"
    r"(?:critically\s+)?"
    r"(?:"
    r"which of the following|"
    r"select|choose|pick|match|complete|"
    r"list|name|identify|define|describe|explain|outline|state|provide|"
    r"illustrate|evaluate|determine|articulate|discuss|analyse|analyze|compare|review|appraise|"
    r"assess|"
    r"what|when|where|why|how|"
    r"must\b"
    r")\b",
    re.IGNORECASE,
)
V3_OPTION_LINE_RE = re.compile(
    r"^\s*(?:"
    r"(?:option\s*\d+)|"
    r"(?:\(?[a-h]\)|[a-h][\.\)])|"
    r"(?:\(?i{1,3}v?\)|i{1,3}v?[\.\)])"
    r")\s+",
    re.IGNORECASE,
)


def v3_looks_like_question_start(text: str) -> bool:
    t = v3_clean_text(text)
    if not t:
        return False
    if t.endswith("?"):
        return True
    if "____" in t or "___" in t:
        return True
    return bool(V3_QUESTION_START_RE.match(t))


def v3_looks_like_answer_guide_bullet(text: str) -> bool:
    t = v3_clean_text(text)
    if not t:
        return False
    if t.lower().startswith(("answer may address", "answer must address", "answer needs to address")):
        return True
    if t.lower() in {"that is blank", "has nothing written in the space provided"}:
        return True
    if t.lower().startswith("does not attempt to answer"):
        return True
    return False


def v3_looks_like_option_line(text: str) -> bool:
    t = v3_clean_text(text)
    if not t:
        return False
    if V3_OPTION_LINE_RE.match(t):
        return True
    if t.lower().startswith(("true", "false")):
        return True
    return False


def v3_is_admin_or_meta_line(text: str) -> bool:
    t = v3_clean_text(text)
    if not t:
        return True
    tl = t.lower()
    if tl.startswith("when you have completed all questions"):
        return True
    if tl.startswith("by submitting your"):
        return True
    if tl.startswith("where a learner is assessed as"):
        return True
    if V3_IGNORE_LINE_RE.match(t) or V3_IGNORE_SECTION_RE.match(t) or V3_IGNORE_TABLE_RE.match(t):
        return True
    if v3_looks_like_answer_guide_bullet(t):
        return True
    if V3_ANSWER_GUIDE_START_RE.match(t):
        return True
    return False


def v3_parse_essay_questions_rule_based(items: list[dict]) -> list[dict]:
    out: list[dict] = []
    seen: set[str] = set()
    in_answer_guide = False
    mcqish_stem_re = re.compile(r"^\s*(select|choose|pick|match|complete)\b", re.IGNORECASE)

    def has_answer_guide_soon(idx: int) -> bool:
        for j in range(idx + 1, min(len(items), idx + 8)):
            t2 = v3_clean_text(items[j].get("text", ""))
            if not t2:
                continue
            if V3_ANSWER_GUIDE_START_RE.match(t2) or V3_ANSWER_GUIDE_ANY_RE.search(t2):
                return True
        return False

    i = 0
    while i < len(items):
        t = v3_clean_text(items[i].get("text", ""))
        if not t or v3_is_admin_or_meta_line(t):
            i += 1
            continue

        if bool(items[i].get("is_red")):
            i += 1
            continue

        if V3_ANSWER_GUIDE_START_RE.match(t) or V3_ANSWER_GUIDE_ANY_RE.search(t):
            in_answer_guide = True
            i += 1
            continue

        if V3_ANSWER_GUIDE_ANY_RE.search(t):
            t = v3_strip_answer_guide(t)
        stem = v3_trim_after_question_mark(v3_strip_q_prefix(t))
        stem = v3_trim_after_sentence_if_long(stem)
        if not stem or len(stem) < 10 or not v3_looks_like_question_start(stem):
            i += 1
            continue
        # Prevent essay-fallback from duplicating MCQs.
        # If the stem looks like an MCQ prompt, do not emit an essay question.
        if "which of the following" in stem.lower() or mcqish_stem_re.match(stem):
            i += 1
            continue

        if in_answer_guide and not (stem.endswith("?") or has_answer_guide_soon(i)):
            i += 1
            continue

        optionish = 0
        for j in range(i + 1, min(len(items), i + 8)):
            t2 = v3_clean_text(items[j].get("text", ""))
            if not t2:
                continue
            if v3_is_admin_or_meta_line(t2):
                continue
            if v3_looks_like_option_line(t2):
                optionish += 1
        if optionish >= 2:
            i += 1
            continue

        if not stem.endswith("?") and not has_answer_guide_soon(i):
            i += 1
            continue

        k = v3_normalize_key(stem)
        if k and k not in seen:
            seen.add(k)
            out.append({"question": stem, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": i, "qnum": None})

        in_answer_guide = has_answer_guide_soon(i)
        i += 1
    return out


def v3_filter_items_for_ai(
    items: list[dict],
    ignore_terms: set[str] | None = None,
    ignore_texts: set[str] | None = None,
    mode: str = "balanced",
) -> list[dict]:
    out: list[dict] = []
    ignore_terms = ignore_terms or set()
    ignore_terms_norm = {v3_normalize_key(t) for t in ignore_terms if v3_normalize_key(t)}
    ignore_term_prefixes = sorted([t for t in ignore_terms_norm if t], key=len, reverse=True)
    ignore_texts = ignore_texts or set()
    ignore_texts_norm = {v3_normalize_key(t) for t in ignore_texts if v3_normalize_key(t)}
    mode = (mode or "balanced").strip().lower()
    if mode not in {"balanced", "loose", "strict"}:
        mode = "balanced"
    in_answer_guide = False
    for it in items:
        t = v3_clean_text(it.get("text", ""))
        if not t:
            continue
        if v3_normalize_key(t) in ignore_texts_norm:
            continue
        if V3_IGNORE_LINE_RE.match(t):
            continue
        if V3_IGNORE_SECTION_RE.match(t):
            continue
        if mode in {"balanced", "strict"}:
            if V3_ANSWER_GUIDE_START_RE.match(t):
                in_answer_guide = True
                continue
            if V3_ANSWER_GUIDE_ANY_RE.search(t):
                pre = v3_strip_answer_guide(t)
                if pre and v3_looks_like_question_start(pre) and len(pre) >= 10:
                    out.append({"text": pre, "is_red": False})
                in_answer_guide = True
                continue
        if V3_IGNORE_TABLE_RE.match(t):
            continue
        if mode in {"balanced", "strict"} and v3_looks_like_answer_guide_bullet(t):
            continue
        tn = v3_normalize_key(t)
        if tn in ignore_terms_norm:
            continue
        skipped = False
        for pref in ignore_term_prefixes:
            if tn.startswith(pref + " "):
                skipped = True
                break
        if skipped:
            continue
        m = re.match(r"^\s*(?:q\s*\d+\s*[\.\)]\s*)?(.*)$", t, flags=re.IGNORECASE)
        if m:
            rest = v3_normalize_key(v3_clean_text(m.group(1)))
            if rest in ignore_terms_norm:
                continue

        if mode == "strict" and in_answer_guide:
            if v3_looks_like_question_start(t):
                in_answer_guide = False
            else:
                if bool(it.get("is_red")):
                    continue
                if len(t) <= 120 and not t.endswith("?"):
                    continue
                continue
        if mode == "balanced" and in_answer_guide:
            if v3_looks_like_question_start(t):
                in_answer_guide = False
            else:
                if len(t) <= 80 and not t.endswith("?"):
                    continue
                if bool(it.get("is_red")) and len(t) <= 80:
                    continue
                continue

        if mode in {"balanced", "strict"} and len(t) <= 20 and V3_COOKERY_METHOD_WORD_RE.match(t):
            continue

        out.append(it)
    return out


def v3_ai_segment_items_openai(items: list[dict], cfg: OpenAIConfig) -> tuple[list[dict], list[str]]:
    log: list[str] = []
    if not items:
        return [], log

    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "questions": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "kind": {"type": "string", "enum": ["mcq", "essay"]},
                        "stem": {"type": "array", "items": {"type": "integer"}},
                        "options": {"type": "array", "items": {"type": "array", "items": {"type": "integer"}}},
                    },
                    "required": ["kind", "stem", "options"],
                },
            }
        },
        "required": ["questions"],
    }

    def to_line(i: int) -> str:
        t = v3_clean_text(items[i].get("text", ""))
        red = "R1" if items[i].get("is_red") else "R0"
        return f"I{i}|{red}|{t}"

    max_block_items = 170
    overlap = 50
    blocks = []
    start = 0
    while start < len(items):
        end = min(len(items), start + max_block_items)
        blocks.append((start, end))
        if end >= len(items):
            break
        start = max(0, end - overlap)

    all_qs: list[dict] = []

    def should_demote_mcq_to_essay(stem_text: str, options: list[str], correct: list[int]) -> bool:
        s = v3_normalize_key(stem_text)
        if not s:
            return False
        if len(options) <= 1:
            return True
        # If we cannot identify any correct options (no red text in any option group),
        # only keep this as an MCQ when the stem clearly signals an MCQ.
        # This prevents short-answer questions that have rubric bullets from being misclassified as MCQ.
        mcq_cue = bool(
            re.search(
                r"\b("
                r"which of the following|"
                r"which strategy or technique|"
                r"stand for|"
                r"select|choose|pick|"
                r"more than one answer|"
                r"select all that apply|choose all that apply"
                r")\b",
                s,
            )
        )
        if not correct and not mcq_cue:
            return True
        if len(options) == 2:
            a, b = v3_normalize_key(options[0]), v3_normalize_key(options[1])
            if a and b and (a in b or b in a):
                return True
        if s.startswith(("what is the name", "what was the name", "what is meant", "what is the origin")):
            if len(options) <= 3 and len(correct) <= 1:
                return True
        for opt in options:
            o = v3_normalize_key(opt)
            if o and len(o) > 25 and (o in s or s in o):
                return True
        return False

    for (a, b) in blocks:
        ctx = []
        for i in range(a, b):
            t = v3_clean_text(items[i].get("text", ""))
            if not t:
                continue
            ctx.append(to_line(i))
        if len(ctx) < 6:
            continue

        log.append(f"AI block: {a}-{b} lines={len(ctx)}")
        prompt = (
            "You are segmenting a DOCX extraction into Canvas quiz questions.\n"
            "Return STRICT JSON only (per schema).\n"
            "\n"
            "Hard rules:\n"
            "- You MUST NOT invent any text.\n"
            "- You may ONLY reference item indices (I<n>) from the provided list.\n"
            "- Keep original order (earlier indices first).\n"
            "- Do NOT create questions from instructions/policy/rubric.\n"
            "- Do NOT include assessor guide content like 'Answer may/must/needs address' or sample answers in stems/options.\n"
            "- For MCQ: include ALL options (both R0 and R1). Mark correctness by whether an option contains any R1 items.\n"
            "- For essay questions: options must be [].\n"
            "\n"
            "MCQ rules:\n"
            "- Options are typically lettered (a), (b), etc or separate lines under a prompt.\n"
            "- If you cannot find at least 2 options, do NOT output an MCQ.\n"
            "\n"
            "Essay rules:\n"
            "- Use the question prompt only.\n"
            "\n"
            "Items (format: I<index>|R0/R1|text):\n"
            + "\n".join(ctx)
        )

        data, err = openai_responses_json_schema(prompt, "segment_questions", schema, cfg)
        if err:
            log.append(f"  block failed: {err}")
            continue
        qs = data.get("questions") if isinstance(data, dict) else None
        if not isinstance(qs, list):
            log.append("  block skipped: missing questions[]")
            continue

        for q in qs:
            if not isinstance(q, dict):
                continue
            kind = (q.get("kind") or "").strip().lower()
            stem_ids = q.get("stem") if isinstance(q.get("stem"), list) else []
            if kind not in ("mcq", "essay"):
                continue
            if not stem_ids or not all(isinstance(x, int) for x in stem_ids):
                continue
            if any(x < 0 or x >= len(items) for x in stem_ids):
                continue

            stem_text = v3_clean_text(" ".join(v3_clean_text(items[x].get("text", "")) for x in stem_ids))
            stem_text = v3_strip_q_prefix(v3_strip_answer_guide(stem_text))
            stem_text = v3_trim_after_question_mark(stem_text)
            stem_text = v3_trim_after_sentence_if_long(stem_text)
            if not stem_text or len(stem_text) < 10:
                continue
            if stem_text.lower().startswith(("answer may address", "answer must address", "answer needs to address")):
                continue
            if not v3_looks_like_question_start(stem_text):
                continue

            if kind == "essay":
                all_qs.append({"question": stem_text, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": min(stem_ids), "qnum": None})
                continue

            opt_groups = q.get("options") if isinstance(q.get("options"), list) else []
            option_texts: list[str] = []
            correct: list[int] = []

            def _looks_like_continuation(text: str) -> bool:
                t2 = (text or "").strip()
                if not t2:
                    return True
                if t2[:1].islower():
                    return True
                if t2.startswith(("•", "-", "–", "—", ",", ";", ":", ")", "]")):
                    return True
                return False

            def _split_group_into_segments(idxs: list[int]) -> list[list[int]]:
                # AI sometimes merges multiple options into one group.
                # Split on any index whose text looks like a new standalone option line.
                if len(idxs) <= 1:
                    return [idxs]
                segments: list[list[int]] = []
                cur: list[int] = []
                for ix in idxs:
                    tx = v3_clean_text(items[ix].get("text", ""))
                    if not cur:
                        cur = [ix]
                        continue
                    if _looks_like_continuation(tx):
                        cur.append(ix)
                        continue
                    segments.append(cur)
                    cur = [ix]
                if cur:
                    segments.append(cur)
                # Avoid pathological splits.
                if len(segments) > 6:
                    return [idxs]
                return segments

            for group in opt_groups:
                if not isinstance(group, list) or not group or not all(isinstance(x, int) for x in group):
                    continue
                if any(x < 0 or x >= len(items) for x in group):
                    continue
                for seg in _split_group_into_segments(group):
                    t = v3_clean_text(" ".join(v3_clean_text(items[x].get("text", "")) for x in seg))
                    if not t:
                        continue
                    if V3_ANSWER_GUIDE_START_RE.match(t) or V3_IGNORE_TABLE_RE.match(t) or V3_IGNORE_LINE_RE.match(t):
                        continue
                    option_texts.append(t)
                    if any(bool(items[x].get("is_red")) for x in seg):
                        correct.append(len(option_texts) - 1)

            seen = set()
            out_opts = []
            out_corr = []
            for i_opt, opt in enumerate(option_texts):
                k = v3_normalize_key(opt)
                if k in seen:
                    continue
                seen.add(k)
                if i_opt in correct:
                    out_corr.append(len(out_opts))
                out_opts.append(opt)
            if len(out_opts) < 2:
                continue

            if should_demote_mcq_to_essay(stem_text, out_opts, out_corr):
                all_qs.append({"question": stem_text, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": min(stem_ids), "qnum": None})
            else:
                all_qs.append({"question": stem_text, "options": out_opts, "correct": out_corr, "multi": ("apply" in stem_text.lower()) or (len(out_corr) > 1), "kind": "mcq", "_order": min(stem_ids), "qnum": None})

    deduped = []
    seen_q = set()
    for q in sorted(all_qs, key=lambda q: int(q.get("_order", 10**9))):
        k = v3_normalize_key(q.get("question", ""))
        if not k or k in seen_q:
            continue
        seen_q.add(k)
        deduped.append(q)
    return deduped, log


def v3_question_dedupe_key(q: dict) -> str:
    kind = (q.get("kind") or "").lower().strip()
    if kind == "matching":
        pairs = q.get("pairs") or []
        parts = []
        for p in pairs:
            left = v3_normalize_key((p or {}).get("left") or "")
            right = v3_normalize_key((p or {}).get("right") or "")
            if left or right:
                parts.append(f"{left}->{right}")
        return "matching|" + v3_normalize_key(q.get("question", "")) + "|" + "|".join(parts)
    if kind == "mcq":
        opts = [v3_normalize_key(o) for o in (q.get("options") or []) if v3_normalize_key(o)]
        return "mcq|" + v3_normalize_key(q.get("question", "")) + "|" + "|".join(opts)
    return "essay|" + v3_normalize_key(q.get("question", ""))


def v3_dedupe_questions(questions: list[dict]) -> tuple[list[dict], int]:
    # Two-pass de-dupe:
    # 1) Remove exact duplicates using a strict key that includes kind/options/pairs.
    # 2) Remove cross-kind duplicates (common in v3) by preferring MCQ over Essay
    #    when the stem text matches.
    strict_kept: list[dict] = []
    strict_seen: set[str] = set()
    removed = 0
    for q in questions:
        k = v3_question_dedupe_key(q)
        if not k or k in strict_seen:
            removed += 1
            continue
        strict_seen.add(k)
        strict_kept.append(q)

    def text_key(q: dict) -> str:
        return v3_normalize_key(v3_clean_text(q.get("question", "")))

    def kind_rank(q: dict) -> int:
        kind = (q.get("kind") or "").lower().strip()
        if kind == "matching":
            return 3
        if kind == "mcq":
            # Prefer MCQs with real options.
            opts = q.get("options") or []
            return 2 if isinstance(opts, list) and len(opts) >= 2 else 1
        return 0  # essay

    best_by_text: dict[str, dict] = {}
    for q in strict_kept:
        tk = text_key(q)
        if not tk:
            continue
        prev = best_by_text.get(tk)
        if prev is None:
            best_by_text[tk] = q
            continue
        if kind_rank(q) > kind_rank(prev):
            best_by_text[tk] = q

    # Preserve original order as much as possible.
    out: list[dict] = []
    used_ids: set[int] = set()
    for q in strict_kept:
        tk = text_key(q)
        pick = best_by_text.get(tk)
        if pick is None:
            continue
        if id(pick) in used_ids:
            if pick is not q:
                removed += 1
            continue
        if pick is not q:
            removed += 1
            continue
        used_ids.add(id(pick))
        out.append(pick)
    return out, removed


@dataclass
class OpenAIConfig:
    api_key: str
    model: str
    base_url: str = "https://api.openai.com"
    timeout_s: int = 120


def openai_responses_json_schema(prompt: str, schema_name: str, schema: dict, cfg: OpenAIConfig) -> tuple[dict | None, str | None]:
    url = cfg.base_url.rstrip("/") + "/v1/responses"
    headers = {"Authorization": f"Bearer {cfg.api_key}", "Content-Type": "application/json"}
    body = {"model": cfg.model, "input": prompt, "text": {"format": {"type": "json_schema", "name": schema_name, "schema": schema, "strict": True}}}
    try:
        r = requests.post(url, headers=headers, json=body, timeout=cfg.timeout_s)
    except Exception as e:
        return None, f"OpenAI request failed: {e}"
    if r.status_code >= 400:
        return None, f"OpenAI error {r.status_code}: {r.text}"
    try:
        data = r.json()
    except Exception as e:
        return None, f"OpenAI JSON parse failed: {e}"
    try:
        out = data["output"][0]["content"][0]
        if out.get("type") == "output_text" and out.get("text"):
            import json

            return json.loads(out["text"]), None
        if out.get("type") == "output_json" and out.get("json"):
            return out["json"], None
        if "text" in out:
            import json

            return json.loads(out["text"]), None
    except Exception as e:
        return None, f"OpenAI response parse failed: {e}"
    return None, "OpenAI returned an unexpected response shape."


def filter_items_for_ai(items: list[dict]) -> list[dict]:
    out = []
    started = False
    for it in items:
        t = clean_text(it.get("text", ""))
        if not t:
            continue
        if not started:
            if is_noise_v2(t):
                continue
            if looks_like_question_start_v2(t):
                started = True
            else:
                continue
        if is_noise_v2(t):
            continue
        if ANSWER_GUIDE_START_RE.match(t):
            continue
        if ANSWER_GUIDE_ANY_RE.search(t):
            pre = strip_answer_guide(t)
            if pre and looks_like_question_start_v2(pre):
                out.append({"text": pre, "is_red": bool(it.get("is_red"))})
            continue
        out.append({"text": t, "is_red": bool(it.get("is_red"))})
    return out


def ai_segment_items(items: list[dict], cfg: OpenAIConfig) -> tuple[list[dict], list[str]]:
    log: list[str] = []
    if not items:
        return [], log

    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "questions": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "kind": {"type": "string", "enum": ["mcq", "essay"]},
                        "stem": {"type": "array", "items": {"type": "integer"}},
                        "options": {"type": "array", "items": {"type": "array", "items": {"type": "integer"}}},
                    },
                    "required": ["kind", "stem", "options"],
                },
            }
        },
        "required": ["questions"],
    }

    def to_line(i: int) -> str:
        t = clean_text(items[i].get("text", ""))
        red = "R1" if items[i].get("is_red") else "R0"
        return f"I{i}|{red}|{t}"

    max_block_items = 170
    overlap = 50
    blocks: list[tuple[int, int]] = []
    start = 0
    while start < len(items):
        end = min(len(items), start + max_block_items)
        blocks.append((start, end))
        if end >= len(items):
            break
        start = max(0, end - overlap)

    all_qs: list[dict] = []
    for (a, b) in blocks:
        ctx = []
        for i in range(a, b):
            t = clean_text(items[i].get("text", ""))
            if not t:
                continue
            ctx.append(to_line(i))
        if len(ctx) < 6:
            continue

        log.append(f"AI block: {a}-{b} lines={len(ctx)}")
        prompt = (
            "You are segmenting a DOCX extraction into Canvas quiz questions.\n"
            "Return STRICT JSON only (per schema).\n"
            "\n"
            "Hard rules:\n"
            "- You MUST NOT invent any text.\n"
            "- You may ONLY reference item indices (I<n>) from the provided list.\n"
            "- Keep original order (earlier indices first).\n"
            "- Do NOT create questions from instructions/policy/rubric.\n"
            "- Do NOT include assessor guide content like 'Answer may/must/needs address' or sample answers in stems/options.\n"
            "- Correct MCQ options are those with R1 (red). For essay questions: options must be [].\n"
            "\n"
            "MCQ rules:\n"
            "- Options are typically lettered (a), (b), etc or separate lines under a prompt.\n"
            "- If you cannot find at least 2 options, do NOT output an MCQ.\n"
            "\n"
            "Essay rules:\n"
            "- Use the question prompt only. Do not append long paragraphs after '?'.\n"
            "\n"
            "Items (format: I<index>|R0/R1|text):\n"
            + "\n".join(ctx)
        )

        data, err = openai_responses_json_schema(prompt, "segment_questions", schema, cfg)
        if err:
            log.append(f"  block failed: {err}")
            continue
        qs = data.get("questions") if isinstance(data, dict) else None
        if not isinstance(qs, list):
            log.append("  block skipped: missing questions[]")
            continue

        for q in qs:
            if not isinstance(q, dict):
                continue
            kind = (q.get("kind") or "").strip().lower()
            stem_ids = q.get("stem") if isinstance(q.get("stem"), list) else []
            if kind not in ("mcq", "essay"):
                continue
            if not stem_ids or not all(isinstance(x, int) for x in stem_ids):
                continue
            if any(x < 0 or x >= len(items) for x in stem_ids):
                continue

            stem_text = clean_text(" ".join(clean_text(items[x].get("text", "")) for x in stem_ids))
            stem_text = strip_q_prefix(strip_answer_guide(stem_text))
            stem_text = trim_after_question_mark(stem_text)
            if not stem_text or len(stem_text) < 10:
                continue
            if stem_text.lower().startswith(("answer may address", "answer must address", "answer needs to address")):
                continue

            if kind == "essay":
                all_qs.append({"question": stem_text, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": min(stem_ids)})
                continue

            opt_groups = q.get("options") if isinstance(q.get("options"), list) else []
            option_texts: list[str] = []
            correct: list[int] = []
            for group in opt_groups:
                if not isinstance(group, list) or not group or not all(isinstance(x, int) for x in group):
                    continue
                if any(x < 0 or x >= len(items) for x in group):
                    continue
                t = clean_text(" ".join(clean_text(items[x].get("text", "")) for x in group))
                if not t:
                    continue
                if ANSWER_GUIDE_START_RE.match(t):
                    continue
                option_texts.append(t)
                if any(bool(items[x].get("is_red")) for x in group):
                    correct.append(len(option_texts) - 1)

            # de-dupe options keep order
            seen = set()
            out_opts = []
            out_corr = []
            for i_opt, opt in enumerate(option_texts):
                k = normalize_key(opt)
                if k in seen:
                    continue
                seen.add(k)
                if i_opt in correct:
                    out_corr.append(len(out_opts))
                out_opts.append(opt)
            if len(out_opts) < 2:
                continue

            all_qs.append(
                {
                    "question": stem_text,
                    "options": out_opts,
                    "correct": out_corr,
                    "multi": ("apply" in stem_text.lower()) or (len(out_corr) > 1),
                    "kind": "mcq",
                    "_order": min(stem_ids),
                }
            )

    all_qs.sort(key=lambda q: int(q.get("_order", 10**9)))
    return all_qs, log


def canvas_headers(canvas_token: str):
    return {"Authorization": f"Bearer {canvas_token}"}


def canvas_whoami(canvas_base_url: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/users/self"
    r = requests.get(url, headers=canvas_headers(canvas_token), timeout=30)
    if r.status_code == 401:
        return None
    r.raise_for_status()
    return r.json()


def list_courses(canvas_base_url: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/courses"
    out = []
    page = 1
    while True:
        r = requests.get(url, headers=canvas_headers(canvas_token), params={"per_page": 100, "page": page}, timeout=60)
        r.raise_for_status()
        batch = r.json()
        if not batch:
            break
        out.extend(batch)
        page += 1
    return out


def get_existing_quiz_titles(canvas_base_url: str, course_id: str, canvas_token: str):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes"
    titles = set()
    page = 1
    while True:
        r = requests.get(url, headers=canvas_headers(canvas_token), params={"page": page, "per_page": 100}, timeout=60)
        r.raise_for_status()
        data = r.json()
        if not data:
            break
        for q in data:
            titles.add((q.get("title") or "").strip())
        page += 1
    return titles


def generate_unique_title(base_title, existing_titles):
    if base_title not in existing_titles:
        return base_title
    i = 1
    while True:
        candidate = f"{base_title} ({i})"
        if candidate not in existing_titles:
            return candidate
        i += 1


def create_canvas_quiz(canvas_base_url: str, course_id: str, canvas_token: str, *, title: str, description_html: str = "", settings: dict | None = None) -> int:
    settings = settings or {}
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes"
    quiz_obj = {
        "title": title,
        "description": description_html,
        "published": False,
        "quiz_type": "assignment",
        "shuffle_answers": bool(settings.get("shuffle_answers", False)),
        "one_question_at_a_time": bool(settings.get("one_question_at_a_time", False)),
        "show_correct_answers": bool(settings.get("show_correct_answers", False)),
        "scoring_policy": settings.get("scoring_policy", "keep_highest"),
    }

    tl = int(settings.get("time_limit", 0) or 0)
    if tl > 0:
        quiz_obj["time_limit"] = tl

    allow_multi = bool(settings.get("allow_multiple_attempts", False))
    if allow_multi:
        aa = int(settings.get("allowed_attempts", 2) or 2)
        quiz_obj["allowed_attempts"] = max(2, aa)

    if bool(settings.get("access_code_enabled", False)) and (settings.get("access_code") or "").strip():
        quiz_obj["access_code"] = settings["access_code"].strip()

    for k in ["due_at", "unlock_at", "lock_at"]:
        v = (settings.get(k) or "").strip()
        if v:
            quiz_obj[k] = v

    payload = {"quiz": quiz_obj}
    r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    if r.status_code == 401:
        raise RuntimeError("401 Unauthorized — token invalid/expired.")
    if r.status_code == 403:
        raise RuntimeError("403 Forbidden — missing permission in this course.")
    r.raise_for_status()
    return r.json()["id"]


def publish_quiz(canvas_base_url: str, course_id: str, canvas_token: str, quiz_id: int):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes/{quiz_id}"
    payload = {"quiz": {"published": True}}
    r = requests.put(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    r.raise_for_status()


def add_question_to_quiz(canvas_base_url: str, course_id: str, canvas_token: str, quiz_id: int, q: dict):
    url = f"{canvas_base_url.rstrip('/')}/courses/{course_id}/quizzes/{quiz_id}/questions"
    qtext = strip_q_prefix((q.get("question") or "").strip())
    kind = (q.get("kind") or "").lower()

    if kind == "matching":
        pairs = q.get("pairs") or []
        answers = []
        for p in pairs:
            left = (p.get("left") or "").strip()
            right = (p.get("right") or "").strip()
            if left and right:
                answers.append({"answer_match_left": left, "answer_match_right": right, "answer_weight": 100})
        payload = {"question": {"question_name": qtext[:100] if qtext else "Matching", "question_text": qtext, "question_type": "matching_question", "points_possible": 1, "answers": answers}}
        r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
        if r.status_code >= 400:
            raise RuntimeError(f"Canvas error {r.status_code}: {r.text[:600]}")
        r.raise_for_status()
        return

    opts = [o.strip() for o in (q.get("options") or []) if o and o.strip()]
    correct = q.get("correct", []) or []
    if kind == "essay" or len(opts) < 2:
        payload = {"question": {"question_name": qtext[:100] if qtext else "Question", "question_text": qtext or " ", "question_type": "essay_question", "points_possible": 1}}
        r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
        r.raise_for_status()
        return

    qlower = (qtext or "").lower()
    multi = bool(q.get("multi")) or (len(correct) > 1) or bool(re.search(r"\bselect\s+(two|three|four|five|\d+)", qlower)) or ("apply" in qlower)
    qtype = "multiple_answers_question" if multi else "multiple_choice_question"
    answers = [{"answer_text": opt, "answer_weight": (100 if idx in correct else 0)} for idx, opt in enumerate(opts)]
    payload = {"question": {"question_name": (qtext[:100] if qtext else "Question"), "question_text": qtext, "question_type": qtype, "points_possible": 1, "answers": answers}}
    r = requests.post(url, headers=canvas_headers(canvas_token), json=payload, timeout=60)
    if r.status_code >= 400:
        raise RuntimeError(f"Canvas error {r.status_code}: {r.text[:600]}")
    r.raise_for_status()


# ===================================================
# Streamlit UI (v1-style flow)
# ===================================================
st.set_page_config(page_title="Canvas Quiz Uploader", layout="wide")
st.title("Canvas Quiz Uploader")



with st.sidebar:
    login_expanded = not bool(st.session_state.logged_in)
    course_expanded = bool(st.session_state.logged_in) and not bool(st.session_state.selected_course_id)
    parser_expanded = True

    with st.expander("🔐 Login", expanded=login_expanded):
        st.session_state.canvas_base_url = st.text_input("Canvas Base URL", value=st.session_state.canvas_base_url).strip()
        st.session_state.canvas_token = st.text_input("Canvas Access Token", value=st.session_state.canvas_token, type="password")
        c_login, c_logout = st.columns(2)
        if c_login.button("Login", use_container_width=True):
            try:
                me = canvas_whoami(st.session_state.canvas_base_url, st.session_state.canvas_token)
                if me:
                    st.session_state.logged_in = True
                    st.session_state.me = me
                    st.session_state.courses_cache = None
                else:
                    st.session_state.logged_in = False
                    st.session_state.me = None
                    st.error("Login failed: token invalid/expired.")
            except Exception as e:
                st.session_state.logged_in = False
                st.session_state.me = None
                st.error(f"Login failed: {e}")

        if c_logout.button("Logout", use_container_width=True):
            st.session_state.logged_in = False
            st.session_state.me = None
            st.session_state.selected_course_id = None
            st.session_state.courses_cache = None
            st.session_state.questions = []
            st.session_state.parsed_ok = False

        if st.session_state.logged_in and st.session_state.me:
            st.caption(f"User: {st.session_state.me.get('name','')}")
        else:
            st.caption("Token login only.")

    with st.expander("✅ Course", expanded=course_expanded):
        if not st.session_state.logged_in:
            st.info("Login first to load courses.")
        else:
            try:
                if st.session_state.courses_cache is None:
                    st.session_state.courses_cache = list_courses(st.session_state.canvas_base_url, st.session_state.canvas_token)
                courses = st.session_state.courses_cache or []
                if not courses:
                    st.warning("No courses visible to this token.")
                else:
                    label_to_id = {}
                    labels = []
                    for c in courses:
                        cid = c.get("id")
                        name = (c.get("name") or c.get("course_code") or f"Course {cid}").strip()
                        label = f"{name} (ID: {cid})"
                        labels.append(label)
                        label_to_id[label] = str(cid)

                    default_index = 0
                    if st.session_state.selected_course_id:
                        for i, lb in enumerate(labels):
                            if label_to_id[lb] == st.session_state.selected_course_id:
                                default_index = i
                                break
                    chosen = st.selectbox("Select course", labels, index=default_index)
                    st.session_state.selected_course_id = label_to_id[chosen]

                    if st.button("Refresh courses", use_container_width=True):
                        st.session_state.courses_cache = None
                        st.rerun()
            except Exception as e:
                st.error(f"Failed to load courses: {e}")

    with st.expander("Parser", expanded=parser_expanded):
        parser_mode = st.selectbox(
            "Version",
            [
                "v1 (rule-based)",
                "v2 (rule-based)",
                "v3 (AI-hybrid)",
            ],
            index=0,
        )
        prev_mode = st.session_state.last_parser_mode
        if prev_mode is None:
            st.session_state.last_parser_mode = parser_mode
        elif prev_mode != parser_mode:
            # Switching parsers should not keep/merge previous parsed results.
            st.session_state.last_parser_mode = parser_mode
            st.session_state.questions = []
            st.session_state.parsed_ok = False
            st.session_state.description_html = ""
            st.session_state.docx_filename = None
            st.session_state.parse_run_id += 1
            # Keep quiz settings, login, and course selection.
            st.rerun()
        if parser_mode.startswith("v3"):
            st.session_state.openai_api_key = st.text_input("OpenAI API key", value=st.session_state.openai_api_key, type="password")
            st.session_state.openai_model = st.text_input("Model", value=st.session_state.openai_model)
            st.session_state.openai_base_url = st.text_input("Base URL", value=st.session_state.openai_base_url)


if not st.session_state.logged_in:
    st.warning("Please login in the sidebar first.")
    st.stop()

if not st.session_state.selected_course_id:
    st.warning("Please select a course in the sidebar.")
    st.stop()

course_id = st.session_state.selected_course_id
canvas_base_url = st.session_state.canvas_base_url
canvas_token = st.session_state.canvas_token


st.subheader("1) Upload DOCX and Parse")
uploaded = st.file_uploader("DOCX", type=["docx"], label_visibility="collapsed")

c_parse1, c_parse2 = st.columns([1, 1])
parse_btn = c_parse1.button("Parse", type="primary", use_container_width=True)
clear_btn = c_parse2.button("Clear parsed results", use_container_width=True)

log_box = st.empty()

V2C_Q_PREFIX_RE = re.compile(
    r"^[^A-Za-z0-9]*(?:lo\s*)?(?:question|q)\s*\d+\s*(?:[\.\)]\s*|[:\-–—]\s+|\s+)",
    re.IGNORECASE,
)
V2C_NUM_PREFIX_RE = re.compile(
    r"^[^A-Za-z0-9]*\(?\d+\)?\s*(?:[\.\)]\s*|[:\-–—]\s+)",
    re.IGNORECASE,
)
V2C_LETTERED_OPT_PREFIX_RE = re.compile(r"^\s*(?:[\(\[]?[a-hA-H][\)\].:-])\s+")


def v2c_clean_text(t: str) -> str:
    return re.sub(r"\s+", " ", (t or "")).strip()


def v2c_normalize_key(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    return s


def v2c_strip_q_prefix(line: str) -> str:
    s = (line or "").strip()
    s = V2C_Q_PREFIX_RE.sub("", s)
    s = V2C_NUM_PREFIX_RE.sub("", s)
    return s.strip()


def v2c_strip_lettered_prefix(t: str) -> str:
    return V2C_LETTERED_OPT_PREFIX_RE.sub("", (t or "").strip()).strip()


V2C_NOISE_RE = re.compile(
    r"^(Instructions|For learners|For students|For assessors|Range and conditions|Decision-making rules|"
    r"Pre-approved reasonable adjustments|Rubric|Knowledge Test|"
    r"A rubric has been assigned\b|Answers will be assessed against\b|As a principle\b)\b",
    re.IGNORECASE,
)
V2C_STOP_OPTION_RE = re.compile(
    r"^(Learner feedback|Assessment outcome|Assessor signature|Assessor name|Final comments)\b"
    r"|^(Competent|Not Yet Competent|NYC|C|Date)\s*[:\-]?$",
    re.IGNORECASE,
)
V2C_OPTION_NOISE_RE = re.compile(r"^(Learning\s+Vault|\d{1,2}/\d{1,2}/\d{2,4}|SIT[A-Z0-9]{5,}\b)", re.IGNORECASE)
V2C_QUESTION_CMD_INNER_RE = re.compile(
    r"\b(Which\s+of\s+the\s+following\b|"
    r"(Identify|Select|Choose|Pick)\s+(?:the\s+)?(one|two|three|four|five|six|seven|eight|nine|ten|\d+)\b)",
    re.IGNORECASE,
)
V2C_COMMAND_QUESTION_RE = re.compile(
    r"^(Illustrate|Explain|Describe|Discuss|Outline|Compare|Summari[sz]e|"
    r"Critically\s+(?:assess|analyse|analyze|evaluate)|"
    r"Evaluate|Determine|Articulate|Prescribe|Analyse|Analyze|Review|Recommend|Provide)\b.+",
    re.IGNORECASE,
)
V2C_RUBRIC_START_RE = re.compile(r"^Answer\s+needs\s+to\s+address\b", re.IGNORECASE)
V2C_ESSAY_GUIDE_RE = re.compile(r"^Answer\s+(may|must)\s+address", re.IGNORECASE)

V2C_MATCHING_STEM_RE = re.compile(
    r"\b("
    r"complete\s+the\s+table|"
    r"drag(?:ging)?\s+and\s+drop(?:ping)?|"
    r"drag\s+and\s+drop|"
    r"match\s+each|"
    r"match\s+the\s+following|"
    r"match\s+.*\s+to\s+the\s+correct|"
    r"select\s+one.*for\s+each"
    r")\b",
    re.IGNORECASE,
)


def v2c_looks_like_matching_stem(t: str) -> bool:
    t2 = v2c_strip_q_prefix(v2c_clean_text(t))
    if not t2:
        return False
    low = t2.lower()
    if low.startswith(("for learners", "for assessors", "for students")):
        return False
    if V2C_COMMAND_QUESTION_RE.match(t2):
        return False
    if "which of the following" in low:
        return False
    return bool(V2C_MATCHING_STEM_RE.search(t2))


V2C_DANGLING_Q_END_RE = re.compile(r"\b(of|for|to|with|and|or|in|on|at|from|by|as|about)\s*$", re.IGNORECASE)


def v2c_merge_dangling_question_lines(items: list[dict]) -> list[dict]:
    out: list[dict] = []
    i = 0
    n = len(items)

    while i < n:
        it = items[i]
        t = v2c_clean_text(it.get("text", ""))
        if not t:
            i += 1
            continue

        t_stem = v2c_strip_q_prefix(t)
        can_start_q = bool(re.match(r"^(?:in\s+)?(which|what|why|how)\b", t_stem, re.IGNORECASE))
        dangling = can_start_q and ("?" not in t_stem) and V2C_DANGLING_Q_END_RE.search(t_stem)
        dangling = bool(dangling) and not v2c_looks_like_matching_stem(t_stem) and not V2C_NOISE_RE.match(t_stem) and not V2C_STOP_OPTION_RE.match(t_stem)

        if dangling and (i + 1) < n:
            nxt = v2c_clean_text(items[i + 1].get("text", ""))
            nxt_stem = v2c_strip_q_prefix(nxt)

            if (
                nxt_stem
                and nxt_stem[:1].islower()
                and not V2C_LETTERED_OPT_PREFIX_RE.match(nxt_stem)
                and not v2c_looks_like_matching_stem(nxt_stem)
                and not V2C_NOISE_RE.match(nxt_stem)
                and not V2C_STOP_OPTION_RE.match(nxt_stem)
            ):
                combined = v2c_clean_text(f"{t_stem} {nxt_stem}")
                if "?" not in combined:
                    combined = combined.rstrip(".") + "?"
                out.append({"text": combined, "is_red": False})
                i += 2
                continue

        out.append(it)
        i += 1

    return out


def v2c_question_fingerprint(q: dict) -> str:
    qt = v2c_normalize_key(q.get("question", ""))
    kind = v2c_normalize_key(q.get("kind", ""))
    opts = [v2c_normalize_key(x) for x in (q.get("options") or [])]
    pairs = q.get("pairs") or []
    pairs_blob = "||".join([v2c_normalize_key((p or {}).get("left", "")) + "=>" + v2c_normalize_key((p or {}).get("right", "")) for p in pairs])
    blob = kind + "||" + qt + "||" + "||".join(opts) + "||" + pairs_blob
    return hashlib.sha1(blob.encode("utf-8")).hexdigest()


def v2c_dedupe_questions(questions: list[dict]) -> list[dict]:
    seen = set()
    out: list[dict] = []
    for q in questions:
        fp = v2c_question_fingerprint(q)
        if fp in seen:
            continue
        seen.add(fp)
        out.append(q)
    return out


def v2c_collapse_duplicate_mcq(questions: list[dict]) -> list[dict]:
    groups: dict[str, list[dict]] = {}
    non_mcq: list[dict] = []

    for q in questions:
        if (q.get("kind") or "").lower() == "mcq":
            key = v2c_normalize_key(q.get("question", ""))
            groups.setdefault(key, []).append(q)
        else:
            non_mcq.append(q)

    def score_mcq(q: dict) -> tuple[int, int]:
        opts = q.get("options") or []
        n = len(opts)
        score = 0
        if 2 <= n <= 6:
            score += 6
        elif 2 <= n <= 10:
            score += 3
        else:
            score -= 3
        if n > 12:
            score -= 8
        bad = 0
        for o in opts:
            low = v2c_normalize_key(o)
            if "?" in (o or ""):
                bad += 1
            if V2C_NOISE_RE.match(o or "") or V2C_OPTION_NOISE_RE.match(o or ""):
                bad += 1
            if low.startswith(("for learners", "for students", "for assessors")):
                bad += 1
        score -= bad * 2
        order = int(q.get("_order", 10**9))
        return (score, -order)

    kept: list[dict] = []
    for key, arr in groups.items():
        arr2 = sorted(arr, key=score_mcq, reverse=True)
        kept.append(arr2[0])

    out = non_mcq + kept
    out.sort(key=lambda q: int(q.get("_order", 10**9)))
    return out


def v2c_parse_mcq_questions(items: list[dict]) -> list[dict]:
    questions_list: list[dict] = []
    current_q: str | None = None
    current_opts: list[dict] = []
    saw_multi_hint = False
    current_start_idx: int | None = None
    pending_multi_hint = False

    instruction_block_re = re.compile(r"^(instructions|for\s+learners|for\s+students|for\s+assessors)\b", re.IGNORECASE)
    meta_line_re = re.compile(r"^(More than one answer may apply|Select all that apply|Choose all that apply)\b", re.IGNORECASE)
    meta_any_re = re.compile(r"\b(More than one answer may apply|Select all that apply|Choose all that apply)\b", re.IGNORECASE)
    colon_stem_re = re.compile(r":\s*(?:\((?:select|choose)\b.*\))?\s*$", re.IGNORECASE)
    question_start_re = re.compile(r"^(?:in\s+)?(which|what|why|how)\b", re.IGNORECASE)
    select_stem_re = re.compile(r"^(?:q\s*\d+\.?\s*)?(select|choose|pick)\s+the\s+(best|correct|most\s+appropriate)\b", re.IGNORECASE)
    read_stem_re = re.compile(r"^(?:q\s*\d+\.?\s*)?read\s+the\s+following\b", re.IGNORECASE)
    complete_stem_re = re.compile(r"^(?:q\s*\d+\.?\s*)?complete\s+the\b", re.IGNORECASE)
    select_hint_re = re.compile(r"\((select|choose)\b", re.IGNORECASE)
    fill_gap_block_re = re.compile(r"\bfill\s+the\s+(gap|blank)\b", re.IGNORECASE)
    contains_select_summary_re = re.compile(r"\b(select|choose|pick)\s+the\s+(best|correct|most\s+appropriate)\s+summary\b", re.IGNORECASE)
    best_match_re = re.compile(r"\b(best\s+match|does\s+the\s+following\s+description\s+best\s+match)\b", re.IGNORECASE)

    def is_strong_stem(txt: str) -> bool:
        t = (txt or "").strip()
        return bool(
            t.endswith("?")
            or V2C_QUESTION_CMD_INNER_RE.search(t)
            or select_hint_re.search(t)
            or meta_any_re.search(t)
            or colon_stem_re.search(t)
            or select_stem_re.match(t)
            or question_start_re.match(t)
        )

    def flush():
        nonlocal current_q, current_opts, saw_multi_hint, current_start_idx
        if not current_q:
            return

        opts = [o for o in current_opts if not V2C_NOISE_RE.match(o["text"])]
        opts = [o for o in opts if not V2C_OPTION_NOISE_RE.match(o["text"])]
        option_texts = [o["text"] for o in opts]
        correct = [i for i, o in enumerate(opts) if o["is_red"]]

        qtext = v2c_strip_q_prefix(current_q.strip())
        qlower = qtext.lower()
        multi = (
            saw_multi_hint
            or bool(re.search(r"\bselect\s+(two|three|four|five|\d+)", qlower))
            or ("apply" in qlower)
            or (len(correct) > 1)
        )

        if len(option_texts) < 2:
            if not is_strong_stem(qtext):
                current_q = None
                current_opts = []
                saw_multi_hint = False
                current_start_idx = None
                return
            option_texts = [
                "⚠ Option text not extracted (likely image/shape). Please replace this option.",
                "⚠ Option text not extracted (likely image/shape). Please replace this option.",
            ]
            correct = []

        questions_list.append({
            "question": qtext,
            "options": option_texts,
            "correct": correct,
            "multi": multi,
            "kind": "mcq",
            "_order": (current_start_idx if current_start_idx is not None else 10**9),
            "qnum": None,
        })

        current_q = None
        current_opts = []
        saw_multi_hint = False
        current_start_idx = None

    def parse_fill_gap_line(line: str):
        if line.count("/") < 2:
            return None
        parts = [p.strip() for p in re.split(r"\s*/\s*", line) if p.strip()]
        if len(parts) < 3:
            return None
        opt0 = parts[0].split()[-1]
        opt_last = parts[-1].split()[0]
        if not opt0 or not opt_last:
            return None
        prefix = parts[0][: -len(opt0)].rstrip()
        suffix = parts[-1][len(opt_last):].lstrip()
        options = [opt0] + parts[1:-1] + [opt_last]
        options = [v2c_clean_text(o) for o in options if v2c_clean_text(o)]
        qtext = v2c_clean_text(f"{prefix} ____ {suffix}".strip())
        if len(qtext) < 10 or len(options) < 3:
            return None
        return qtext, options

    def has_plausible_options(start_idx: int) -> bool:
        n = len(items)
        count = 0
        for j in range(start_idx, min(n, start_idx + 25)):
            raw = v2c_clean_text(items[j].get("text", ""))
            if not raw or V2C_NOISE_RE.match(raw):
                continue
            t = v2c_strip_q_prefix(raw)
            if V2C_ESSAY_GUIDE_RE.match(t) or V2C_RUBRIC_START_RE.match(t):
                return False
            if v2c_looks_like_matching_stem(t):
                return False
            if select_stem_re.match(t) or V2C_QUESTION_CMD_INNER_RE.search(t) or best_match_re.search(t) or contains_select_summary_re.search(t):
                break
            if select_hint_re.search(t):
                break
            if t.endswith("?") and len(t) >= 10:
                break
            if len(t) <= 200 and not t.endswith("?"):
                count += 1
                if count >= 2:
                    return True
        return False

    for idx, it in enumerate(items):
        line = v2c_clean_text(it.get("text", ""))
        if not line:
            continue
        if V2C_NOISE_RE.match(line):
            continue
        if V2C_OPTION_NOISE_RE.match(line):
            continue

        if instruction_block_re.match(v2c_strip_q_prefix(line)):
            flush()
            current_q = None
            current_opts = []
            saw_multi_hint = False
            current_start_idx = None
            pending_multi_hint = False
            continue
        if V2C_ESSAY_GUIDE_RE.match(line):
            current_q = None
            current_opts = []
            continue
        if current_q and V2C_STOP_OPTION_RE.match(line):
            flush()
            current_q = None
            current_opts = []
            continue

        t_stem = v2c_strip_q_prefix(line)

        if current_q is None and meta_line_re.match(t_stem):
            pending_multi_hint = True
            continue
        if v2c_looks_like_matching_stem(t_stem):
            flush()
            current_q = None
            current_opts = []
            saw_multi_hint = False
            current_start_idx = None
            continue

        if (
            (current_q is None or len(current_opts) >= 2)
            and colon_stem_re.search(t_stem)
            and len(t_stem) >= 12
            and not v2c_looks_like_matching_stem(t_stem)
            and not V2C_COMMAND_QUESTION_RE.match(t_stem)
            and not V2C_STOP_OPTION_RE.match(line)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = bool(meta_any_re.search(t_stem) or pending_multi_hint)
            current_start_idx = idx
            pending_multi_hint = False
            continue

        if (
            (current_q is None or len(current_opts) >= 2)
            and meta_any_re.search(t_stem)
            and not meta_line_re.match(t_stem)
            and len(t_stem) >= 12
            and not v2c_looks_like_matching_stem(t_stem)
            and not V2C_COMMAND_QUESTION_RE.match(t_stem)
            and not V2C_STOP_OPTION_RE.match(line)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = True
            current_start_idx = idx
            pending_multi_hint = False
            continue

        if fill_gap_block_re.search(t_stem):
            flush()
            current_q = None
            current_opts = []
            saw_multi_hint = False
            current_start_idx = None
            continue

        if (current_q is None) and ("/" in t_stem) and has_plausible_options(idx + 1):
            parsed = parse_fill_gap_line(t_stem)
            if parsed:
                qtext, opts = parsed
                questions_list.append({"question": qtext, "options": opts, "correct": [], "multi": False, "kind": "mcq", "_order": idx, "qnum": None})
                continue

        if select_hint_re.search(t_stem) and (current_q is None or len(current_opts) >= 2) and has_plausible_options(idx + 1):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        if (
            (read_stem_re.match(t_stem) or complete_stem_re.match(t_stem))
            and ("select" in t_stem.lower() or "most appropriate" in t_stem.lower() or "complete" in t_stem.lower())
            and not v2c_looks_like_matching_stem(t_stem)
            and (current_q is None or len(current_opts) >= 2)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        if select_stem_re.match(line) and not v2c_looks_like_matching_stem(line) and not meta_line_re.match(line):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        m = V2C_QUESTION_CMD_INNER_RE.search(line)
        if m:
            flush()
            start = m.start()
            stem = line[:start].strip()
            cmd_plus = line[start:].strip()
            q_line = f"{stem} {cmd_plus}".strip() if stem else cmd_plus
            current_q = v2c_strip_q_prefix(q_line)
            current_opts = []
            current_start_idx = idx
            saw_multi_hint = pending_multi_hint
            pending_multi_hint = False
            continue

        if current_q and not current_opts:
            if (
                current_q and not current_q.strip().endswith("?")
                and line[:1].islower()
                and not V2C_QUESTION_CMD_INNER_RE.search(line)
                and not v2c_looks_like_matching_stem(line)
                and not V2C_COMMAND_QUESTION_RE.match(v2c_strip_q_prefix(line))
                and not V2C_STOP_OPTION_RE.match(line)
                and not meta_line_re.match(line)
            ):
                current_q = (current_q + " " + line).strip()
                continue

        if (
            question_start_re.match(t_stem)
            and len(t_stem) >= 12
            and not v2c_looks_like_matching_stem(t_stem)
            and not V2C_COMMAND_QUESTION_RE.match(t_stem)
            and not V2C_STOP_OPTION_RE.match(line)
            and not meta_line_re.match(line)
            and (
                "?" in t_stem
                or best_match_re.search(t_stem)
                or contains_select_summary_re.search(t_stem)
                or re.search(r"\((select|choose)\b", t_stem, re.IGNORECASE)
            )
            and (current_q is None or len(current_opts) >= 2)
            and has_plausible_options(idx + 1)
        ):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        if (contains_select_summary_re.search(t_stem) or best_match_re.search(t_stem)) and (current_q is None or len(current_opts) >= 2) and has_plausible_options(idx + 1):
            flush()
            current_q = t_stem
            current_opts = []
            saw_multi_hint = pending_multi_hint
            current_start_idx = idx
            pending_multi_hint = False
            continue

        if current_q and meta_line_re.match(line):
            saw_multi_hint = True
            continue

        t = v2c_strip_q_prefix(line)
        if t.endswith("?") and len(t) >= 10 and not V2C_COMMAND_QUESTION_RE.match(t) and not v2c_looks_like_matching_stem(t) and has_plausible_options(idx + 1):
            flush()
            current_q = t
            current_opts = []
            current_start_idx = idx
            continue

        if current_q:
            current_opts.append({"text": line, "is_red": bool(it.get("is_red", False))})

    flush()
    return [q for q in questions_list if len(q.get("options") or []) >= 2 and len(q.get("question") or "") >= 10]


def v2c_parse_essay_questions(items: list[dict]) -> list[dict]:
    questions: list[dict] = []
    n = len(items)
    i = 0
    while i < n:
        raw = v2c_clean_text(items[i].get("text", ""))
        if not raw or V2C_NOISE_RE.match(raw):
            i += 1
            continue
        line = v2c_strip_q_prefix(raw)
        if V2C_COMMAND_QUESTION_RE.match(line):
            j = i + 1
            next_line = ""
            while j < n:
                nxt = v2c_clean_text(items[j].get("text", ""))
                if nxt and not V2C_NOISE_RE.match(nxt):
                    next_line = nxt
                    break
                j += 1
            if V2C_RUBRIC_START_RE.match(next_line):
                questions.append({"question": line, "options": [], "correct": [], "multi": False, "kind": "essay", "_order": i, "qnum": None})
                i = j + 1
                continue
        i += 1
    return [q for q in questions if len((q.get("question") or "").strip()) >= 10]

if clear_btn:
    st.session_state.questions = []
    st.session_state.parsed_ok = False
    st.session_state.description_html = ""
    st.session_state.docx_filename = None
    st.session_state.parse_run_id += 1
    st.rerun()


if parse_btn:
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        if not uploaded:
            raise RuntimeError("Upload a DOCX to parse.")

        uploaded_name = (uploaded.name or "").strip() or "Quiz.docx"
        uploaded_bytes = uploaded.getvalue()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_bytes)
            docx_path = tmp.name

        st.session_state.docx_filename = uploaded_name

        # Always build description from v1 extraction (stable, matches app (1).py).
        desc_items = extract_items_with_red_v1(docx_path)
        st.session_state.description_html = build_description_v1(desc_items)

        qs: list[dict] = []
        ai_log: list[str] = []
        removed_dupes = 0

        if parser_mode.startswith("v1"):
            items_v1 = desc_items
            matching = parse_matching_questions_doc_order_v1_exact(docx_path)
            mcq = parse_mcq_questions_v1(items_v1)
            essay = parse_essay_questions_v1(items_v1)
            qs = dedupe_questions(matching + mcq + essay)

        elif parser_mode.startswith("v2"):
            # v2 (rule-based) — intended to match the user's canvascersion2.py.codex_modified_2026-02-05.py behavior.
            # We reuse the v1 parser, but also split single extracted lines that contain multiple embedded "Q<n>."
            # blocks. This reduces the chance of questions/options being merged.
            items_v2 = v2c_merge_dangling_question_lines(desc_items)
            items_v2 = v2_split_items_on_internal_qnums(items_v2)
            matching = parse_matching_questions_doc_order_v1_exact(docx_path)
            mcq = v2c_parse_mcq_questions(items_v2)
            essay = v2c_parse_essay_questions(items_v2)
            qs = v2c_dedupe_questions(matching + mcq + essay)
            qs = v2c_collapse_duplicate_mcq(qs)

        else:
            # v3 (AI+fallback) defaults mirror ai 43.py
            if not (st.session_state.openai_api_key or "").strip():
                raise RuntimeError("v3 (AI+fallback) requires an OpenAI API key.")
            cfg = OpenAIConfig(
                api_key=st.session_state.openai_api_key.strip(),
                model=(st.session_state.openai_model or "gpt-4.1-mini").strip(),
                base_url=(st.session_state.openai_base_url or "https://api.openai.com").strip(),
            )
            ignore_tables = False
            include_table_essays = True
            include_rule_essay_fallback = True
            ai_filter_mode = "balanced"

            items = v3_extract_items_with_red(docx_path, include_tables=(not ignore_tables))
            items = v3_split_items_on_internal_qnums(items)

            matching = [] if ignore_tables else v3_parse_matching_questions_doc_order(docx_path, items)

            table_essays: list[dict] = []
            if include_table_essays and not ignore_tables:
                table_essays.extend(v3_parse_table_defined_terms_as_essays(docx_path, items))
                table_essays.extend(v3_parse_table_characteristics_as_essays(docx_path, items))

            ignore_terms: set[str] = set()
            for q in table_essays:
                qt = v3_clean_text(q.get("question", ""))
                m = re.match(r"^Define:\s*(.+?)\s*\.", qt, flags=re.IGNORECASE)
                if m:
                    ignore_terms.add(v3_clean_text(m.group(1)))
                    continue
                m = re.match(r"^Describe the essential characteristics of:\s*(.+?)\s*\.", qt, flags=re.IGNORECASE)
                if m:
                    ignore_terms.add(v3_clean_text(m.group(1)))

            ignore_texts: set[str] = set()
            if not ignore_tables:
                ignore_texts = v3_collect_ignore_texts_from_forced_tables(docx_path)

            ai_qs = []
            if (st.session_state.openai_api_key or "").strip():
                ai_input = v3_filter_items_for_ai(items, ignore_terms=ignore_terms, ignore_texts=ignore_texts, mode=ai_filter_mode)
                ai_qs, ai_log = v3_ai_segment_items_openai(ai_input, cfg)

            rule_essays: list[dict] = []
            if include_rule_essay_fallback:
                rule_essays = v3_parse_essay_questions_rule_based(items)

            qs = []
            qs.extend(matching)
            qs.extend(table_essays)
            qs.extend(ai_qs)
            qs.extend(rule_essays)
            qs.sort(key=lambda q: int(q.get("_order", 10**9)))
            qs, removed_dupes = v3_dedupe_questions(qs)

        st.session_state.questions = qs
        st.session_state.parsed_ok = True
        st.session_state.parse_run_id += 1

        # Reset title/instructions defaults for the newly parsed document.
        # (Streamlit widgets keep old values unless their keys change.)
        st.session_state.details["quiz_title"] = ""
        st.session_state.details["quiz_instructions"] = ""

        print("DEBUG: parser:", parser_mode)
        print("DEBUG: items extracted (description v1):", len(desc_items))
        print("DEBUG: matching:", sum(1 for q in qs if (q.get("kind") == "matching")))
        print("DEBUG: mcq:", sum(1 for q in qs if (q.get("kind") == "mcq")))
        print("DEBUG: essay:", sum(1 for q in qs if (q.get("kind") == "essay")))
        if parser_mode.startswith("v3"):
            print("DEBUG: removed_dupes:", removed_dupes)
        print("Parsed questions:", len(qs))
        for ln in ai_log[:60]:
            print(ln)

    log_box.code(buf.getvalue())
    st.success(f"✅ Parsed {len(st.session_state.questions)} questions.")


questions = st.session_state.questions or []
if not questions:
    st.stop()


st.subheader("2) Details (Canvas Quiz Settings)")
default_title = os.path.splitext(st.session_state.docx_filename or "Quiz")[0]
d = st.session_state.details
run = st.session_state.parse_run_id
quiz_title = st.text_input("Quiz Title *", value=(d.get("quiz_title") or default_title), key=f"{run}_quiz_title")
quiz_instructions = st.text_area(
    "Quiz Instructions (HTML allowed)",
    value=(d.get("quiz_instructions") or st.session_state.description_html or ""),
    height=180,
    key=f"{run}_quiz_instructions",
)

d["quiz_title"] = quiz_title
d["quiz_instructions"] = quiz_instructions

c1, c2, c3 = st.columns(3)
d["shuffle_answers"] = c1.checkbox("Shuffle Answers", value=bool(d.get("shuffle_answers", True)))
d["one_question_at_a_time"] = c2.checkbox("Show one question at a time", value=bool(d.get("one_question_at_a_time", False)))
d["show_correct_answers"] = c3.checkbox("Let Students See The Correct Answers", value=bool(d.get("show_correct_answers", False)))

c4, c5, c6 = st.columns(3)
d["time_limit"] = c4.number_input("Time Limit (minutes, 0 = none)", min_value=0, max_value=1440, value=int(d.get("time_limit", 0) or 0), step=5)
d["allow_multiple_attempts"] = c5.checkbox("Allow Multiple Attempts", value=bool(d.get("allow_multiple_attempts", False)))
if d["allow_multiple_attempts"]:
    cur_attempts = int(d.get("allowed_attempts", 2) or 2)
    if cur_attempts < 2:
        cur_attempts = 2
    d["allowed_attempts"] = c5.number_input("Allowed Attempts", min_value=2, max_value=20, value=cur_attempts, step=1)
else:
    d["allowed_attempts"] = 1
d["scoring_policy"] = c6.selectbox("Quiz Score to Keep", ["keep_highest", "keep_latest"], index=0 if d.get("scoring_policy", "keep_highest") == "keep_highest" else 1)

st.markdown("**Quiz Restrictions**")
d["access_code_enabled"] = st.checkbox("Require an access code", value=bool(d.get("access_code_enabled", False)))
if d["access_code_enabled"]:
    d["access_code"] = st.text_input("Access code", value=(d.get("access_code") or ""))
else:
    d["access_code"] = ""

st.markdown("**Assign / Availability (optional, ISO datetime)**")
st.caption("Example: 2026-01-20T23:59:00Z (If you don’t know, leave blank for now.)")
cc1, cc2, cc3 = st.columns(3)
d["due_at"] = cc1.text_input("Due Date (due_at)", value=(d.get("due_at") or ""))
d["unlock_at"] = cc2.text_input("Available from (unlock_at)", value=(d.get("unlock_at") or ""))
d["lock_at"] = cc3.text_input("Until (lock_at)", value=(d.get("lock_at") or ""))

st.session_state.details = d


# ===================================================
# Questions editor (v1-style)
# ===================================================
st.divider()
st.subheader("3) Questions")

page_size = int(st.session_state.get("questions_page_size") or 10)
if page_size not in {5, 10, 15, 20, 30}:
    page_size = 10
total = len(questions)
total_pages = max(1, math.ceil(total / page_size))
page = int(st.session_state.get("questions_page") or 1)
if page < 1:
    page = 1
if page > total_pages:
    page = total_pages

start = (page - 1) * page_size
end = min(start + page_size, total)
st.caption(f"Showing questions {start+1}–{end} of {total}")

edited = [q.copy() for q in questions]
run = st.session_state.parse_run_id

for i in range(start, end):
    q = edited[i]
    kind = (q.get("kind") or "").lower()
    preview = strip_q_prefix(q.get("question", ""))[:90]
    label_kind = "Matching" if kind == "matching" else ("Essay/Short Answer" if kind == "essay" else "MCQ")

    with st.expander(f"Q{i+1} ({label_kind}): {preview}"):
        q_text = st.text_area("Question text", value=q.get("question", ""), key=f"{run}_qtext_{i}", height=90)
        q["question"] = strip_q_prefix(q_text.strip())

        if kind == "essay":
            st.info("This question will be uploaded as an ESSAY (student types the answer).")
            q["options"] = []
            q["correct"] = []
            q["multi"] = False

        elif kind == "matching":
            st.info("This question will be uploaded as MATCHING (left item → dropdown right item).")
            st.caption("Tip: if right side was a bullet list in Word, it appears joined with '; ' — that is correct.")
            pairs = q.get("pairs") or []
            new_pairs = []
            for j, p in enumerate(pairs):
                lc1, lc2 = st.columns([0.6, 0.4])
                left = lc1.text_input(f"Left (row {j+1})", value=p.get("left", ""), key=f"{run}_match_{i}_l_{j}")
                right = lc2.text_input(f"Right (row {j+1})", value=p.get("right", ""), key=f"{run}_match_{i}_r_{j}")
                if left.strip() and right.strip():
                    new_pairs.append({"left": left.strip(), "right": right.strip()})
            q["pairs"] = new_pairs

        else:
            opts = q.get("options", []) or []
            correct_set = set(q.get("correct", []) or [])
            st.write("**Options** (tick ✅ for correct answer)")
            new_opts = []
            new_correct = []
            for j, opt in enumerate(opts):
                oc1, oc2 = st.columns([0.12, 0.88])
                is_corr = oc1.checkbox("", value=(j in correct_set), key=f"{run}_q{i}_corr_{j}")
                opt_text = oc2.text_input(f"Option {j+1}", value=opt, key=f"{run}_q{i}_opt_{j}")
                new_opts.append(opt_text.strip())
                if is_corr:
                    new_correct.append(j)

            add_opt = st.text_input("New option text (optional)", value="", key=f"{run}_q{i}_newopt")
            if add_opt.strip():
                new_opts.append(add_opt.strip())

            cleaned_opts = []
            idx_map = {}
            for old_index, txt in enumerate(new_opts):
                if txt.strip():
                    idx_map[old_index] = len(cleaned_opts)
                    cleaned_opts.append(txt.strip())

            remapped_correct = []
            for old_i in new_correct:
                if old_i in idx_map:
                    remapped_correct.append(idx_map[old_i])

            q["options"] = cleaned_opts
            q["correct"] = sorted(set(remapped_correct))
            qlower = (q.get("question") or "").lower()
            q["multi"] = (("apply" in qlower) or (len(q["correct"]) > 1) or bool(re.search(r"\bselect\s+(two|three|four|five|\d+)", qlower)))

edited = dedupe_questions(edited)
st.session_state.questions = edited

# Pager (bottom): avoids needing to scroll back up after reviewing questions.
st.divider()
pc1, pc2 = st.columns([0.55, 0.45])
new_page_size = pc1.selectbox(
    "Questions per page",
    [5, 10, 15, 20, 30],
    index=[5, 10, 15, 20, 30].index(page_size),
)
new_total_pages = max(1, math.ceil(total / int(new_page_size)))

if new_total_pages <= 200:
    new_page = pc2.selectbox(
        "Page",
        list(range(1, new_total_pages + 1)),
        index=max(0, min(page, new_total_pages) - 1),
        format_func=lambda n: f"{n} / {new_total_pages}",
    )
else:
    new_page = pc2.number_input(
        "Page",
        min_value=1,
        max_value=new_total_pages,
        value=min(page, new_total_pages),
        step=1,
    )

st.session_state.questions_page_size = int(new_page_size)
st.session_state.questions_page = int(new_page)

st.divider()
st.subheader("4) Save to Canvas")

colS1, colS2 = st.columns([1, 1])
save_draft = colS1.button("💾 Save to Canvas (Draft)")
save_publish = colS2.button("🚀 Save & Publish")


def validate_before_upload(qs: list[dict]) -> list[str]:
    problems = []
    for idx, q in enumerate(qs, start=1):
        kind = (q.get("kind") or "").lower()
        qt = (q.get("question") or "").strip()
        if len(qt) < 10:
            problems.append(f"Q{idx}: question text too short.")
        if kind == "matching":
            pairs = q.get("pairs") or []
            if len(pairs) < 2:
                problems.append(f"Q{idx}: matching needs at least 2 pairs.")
            continue
        if kind == "essay":
            continue
        opts = q.get("options") or []
        corr = q.get("correct") or []
        if len(opts) >= 2 and len(corr) == 0:
            problems.append(f"Q{idx}: no correct answer selected (red not detected or tick ✅).")
    return problems


if save_draft or save_publish:
    qs = st.session_state.questions or []
    probs = validate_before_upload(qs)
    if probs:
        st.error("Please fix these issues before uploading:")
        for p in probs[:15]:
            st.write(f"- {p}")
        st.stop()

    base_title = (quiz_title or "").strip() or default_title
    try:
        existing_titles = get_existing_quiz_titles(canvas_base_url, course_id, canvas_token)
        final_title = generate_unique_title(base_title, existing_titles)

        with st.spinner("Creating quiz in Canvas..."):
            quiz_id = create_canvas_quiz(
                canvas_base_url=canvas_base_url,
                course_id=course_id,
                canvas_token=canvas_token,
                title=final_title,
                description_html=quiz_instructions,
                settings=st.session_state.details,
            )

        with st.spinner("Uploading questions..."):
            for q in qs:
                add_question_to_quiz(canvas_base_url, course_id, canvas_token, quiz_id, q)

        if save_publish:
            with st.spinner("Publishing quiz..."):
                publish_quiz(canvas_base_url, course_id, canvas_token, quiz_id)

        st.success("✅ Done!")
        st.write(f"**Quiz title:** {final_title}")
        st.write(f"**Quiz ID:** {quiz_id}")
        st.write(f"**Course ID:** {course_id}")
        st.info("Quiz published ✅" if save_publish else "Quiz saved as draft (unpublished).")
    except Exception as e:
        st.error(f"❌ Upload failed: {e}")

st.caption("Token login only (Canvas API does not support username/password).")
